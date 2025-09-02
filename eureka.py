### Step 1: Importing necessary libraries
import streamlit as st
import requests
import re
import time
import qrcode
from telegraph import Telegraph
import google.generativeai as genai
from bs4 import BeautifulSoup
from io import BytesIO
from deep_translator import GoogleTranslator
from streamlit_javascript import st_javascript
from user_agents import parse
import PIL.Image
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfReader
import firebase_admin
from firebase_admin import credentials, db
from streamlit_chat import message
from streamlit_extras.stylable_container import stylable_container
import toml
from pathlib import Path
import uuid
from text_translation_data import languages_dict, text_translation_data
import copy
import base64
import mimetypes
from serpapi import GoogleSearch
import markdown
import datetime
import random
import json
import streamlit.components.v1 as components
import traceback

# Step 2: Setting Streamlit's page configuration, functions, APIs, database, and session states

# Text Translation
if "language" not in st.session_state:
	st.session_state.language = "english"

def translate_text(text: str, code: int) -> str:
	"""
	Translates the given text into the specified user's language.

	Args:
		text (str): The text to be translated.
		code (int): The language code to translate the text into.

	Returns:
		str: The translated text.
	"""
	try:
		language = st.session_state.language
		translated = text_translation_data[code][language.lower()]
		return translated
	except:
		return text
	

def google_translate(text: str, language: str) -> str:
	try:
		if not language:
			language = st.session_state.language

		if language.lower().strip() == "chinese":
			language = "chinese (simplified)"

		target_lang = languages_dict[language.lower()]
		translated = GoogleTranslator(
			source="en", target=target_lang
		).translate(text)
		return translated
	except:
		return text
	

def is_rtl(language: str):
	if language.lower().strip() == "chinese":
		language = "chinese (simplified)"

	rtl_langs = ["ar", "he", "fa", "ur", "ps", "ku", "yi"]
	lang_code = languages_dict[language.lower()]

	if lang_code in rtl_langs:
		return True
	return False
	

# Setting Streamlit's page configuration
st.set_page_config(
	page_title="Eureka",
	page_icon="assets/favicon.png",
	layout="wide",
	initial_sidebar_state="auto",
)

# Session States
# App's Session States
if "current_page" not in st.session_state:
	st.session_state.current_page = "🧠 Ask AI"

# Login/Signup Session States
if "login_or_signup" not in st.session_state:
	st.session_state.login_or_signup = "login"
if "logged_in" not in st.session_state:
	st.session_state.logged_in = False
if "username" not in st.session_state:
	st.session_state.username = None
if "user_real_name" not in st.session_state:
	st.session_state.user_real_name = None
if "age" not in st.session_state:
	st.session_state.age = None

# Settings Session States
if "theme" not in st.session_state:
	st.session_state.theme = "light"

# Ask AI session states
if "current_chat_num" not in st.session_state:
	st.session_state.current_chat_num = None
if "new_chat" not in st.session_state:
	st.session_state.new_chat = True
if "chat_history" not in st.session_state:
	st.session_state.chat_history = []
if "user_can_send" not in st.session_state:
	st.session_state.user_can_send = True

# Learning Plans and Learning Dashboard Session States
if "make_new_plan" not in st.session_state:
	st.session_state.make_new_plan = False
if "start_learning" not in st.session_state:
	st.session_state.start_learning = False
if "timer_running" not in st.session_state:
	st.session_state.timer_running = False
if "content_marked_done" not in st.session_state:
	st.session_state.content_marked_done = False
if "video_marked_done" not in st.session_state:
	st.session_state.video_marked_done = False
if "flashcard_marked_done" not in st.session_state:
	st.session_state.flashcard_marked_done = False
if "start_quiz" not in st.session_state:
	st.session_state.start_quiz = False
if "quiz_taken" not in st.session_state:
	st.session_state.quiz_taken = False
if "quiz_taken_data" not in st.session_state:
	st.session_state.quiz_taken_data = []
	
# AR, Content and Quiz Session States
if "user_content_length" not in st.session_state:
	st.session_state.user_content_length = None
if "user_content_complexity" not in st.session_state:
	st.session_state.user_content_complexity = None
if "user_custom_prompt" not in st.session_state:
	st.session_state.user_custom_prompt = None
if "user_grade" not in st.session_state:
	st.session_state.user_grade = None

if "target_language" not in st.session_state:
	st.session_state.target_language = "en"

# History page session states
if "history_model_names" not in st.session_state:
	st.session_state.history_model_names = []
if "history_model_thumbs" not in st.session_state:
	st.session_state.history_model_thumbs = []

if "history_ar_links" not in st.session_state:
	st.session_state.history_ar_links = []
if "history_ar_qrs" not in st.session_state:
	st.session_state.history_ar_qrs = []

if "history_content_links" not in st.session_state:
	st.session_state.history_content_links = []
if "history_content_qrs" not in st.session_state:
	st.session_state.history_content_qrs = []
if "history_content_docs" not in st.session_state:
	st.session_state.history_content_docs = []

if "history_quiz_links" not in st.session_state:
	st.session_state.history_quiz_links = []
if "history_quiz_qrs" not in st.session_state:
	st.session_state.history_quiz_qrs = []
if "history_quiz_docs" not in st.session_state:
	st.session_state.history_quiz_docs = []

if "backend_running" not in st.session_state:
	st.session_state.backend_running = False
if "history_rerun" not in st.session_state:
	st.session_state.history_rerun = True
if "history_rerun2" not in st.session_state:
	st.session_state.history_rerun2 = True

# Configuring and Setting up Gemini API
GOOGLE_API_KEY = st.secrets['GOOGLE_API_KEY']
genai.configure(api_key=GOOGLE_API_KEY)

if "gemini_2_0_flash" not in st.session_state:
	st.session_state.gemini_2_0_flash = genai.GenerativeModel("gemini-2.0-flash")
if "gemini_2_5_flash" not in st.session_state:
	st.session_state.gemini_2_5_flash = genai.GenerativeModel("gemini-2.5-flash")
if "gemini_lite_model" not in st.session_state:
	st.session_state.gemini_2_5_flash_lite = genai.GenerativeModel("gemini-2.5-flash-lite")

# Detecting the user's device type to customize the layout
if "device_type" not in st.session_state:
	st.session_state.device_type = None
	ua = st_javascript("window.navigator.userAgent;")

	if isinstance(ua, str) and ua:
		st.session_state.is_pc = parse(ua).is_pc
	else:
		pass

# Setting Up the Firebase Realtime Database
service_account_key_path = "eureka-8bf2e-firebase-adminsdk-fbsvc-5564c98f4a.json"
try:
	if not firebase_admin._apps:  # Check if no default app is already initialized
		cred = credentials.Certificate(service_account_key_path)
		firebase_admin.initialize_app(
			cred,
			{
				"databaseURL": "https://eureka-8bf2e-default-rtdb.europe-west1.firebasedatabase.app"
			},
		)

except Exception as e:
	pass

# Database functions and default folders
root_ref = db.reference("/")
users_ref = root_ref.child("users")
app_data_ref = root_ref.child("app_data")
default_data_ref = app_data_ref.child("default_data")

if users_ref.get() is None:
	users_ref.set({})

if app_data_ref.get() is None:
	app_data_ref.set({})

default_data_ref.update(
	{
		"themes": {
			"light": {
				"primaryColor": "#ff4b4b",
				"backgroundColor": "#ffffff",
				"secondaryBackgroundColor": "#f0f2f6",
				"textColor": "#31333F",
			},
			"dark": {
				"primaryColor": "#ff4b4b",
				"backgroundColor": "#0e1117",
				"secondaryBackgroundColor": "#262730",
				"textColor": "#fafafa",
			},
		}
	}
)


def check_username_validity(username: str) -> bool:
	"""
	Check if a username is valid. A valid username should only contain letters, numbers, underscores, and hyphens.

	Args:
		username (str): The username to check.

	Returns:
		bool: True if the username is valid, False otherwise.
	"""

	allowed = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789-_"
	for char in username:
		if char not in allowed:
			return False  # Found invalid character
	return True  # All characters are valid


def check_password_validity(password: str) -> bool:
	"""
	Check if a password is valid. A valid password should be at least 4 characters long.

	Args:
		password (str): The password to check.

	Returns:
		bool: True if the password is valid, False otherwise.
	"""

	return len(password) >= 4


def register_user(username: str, password: str, name: str, age: int) -> bool:
	"""
	Registers a new user in the Firebase Realtime Database.

	This function attempts to add a new user with the provided details to the
	'users' reference in Firebase. It first checks if a user with the given
	username already exists. If the username is unique, a new user entry is
	created with their password, personal information (name and age), and
	default history and settings.

	Args:
		username (str): The unique username for the new user.
		password (str): The password for the new user.
		name (str): The real name of the new user.
		age (Union[int, str]): The age of the new user. Can be an integer or a string.

	Returns:
		bool: True if the user was successfully registered (username was unique),
		False otherwise (username already exists).
	"""

	current_users = users_ref.get() or {}

	if username not in current_users:
		users_ref.child(username).set(
			{
				"password": password,
				"info": {"name": name.title(), "age": age},
				"history": {
					"ai": {},  # ex: {'chat1': {'title': 'chat title', 'data': [{'role': 'user', 'parts': [{'text': 'Hello Gemini, I am Ibrahim'}]}}, {'role': 'assistant', 'parts': [{'text': 'Hello Ibrahim, I am Gemini'}]}]}
					"total_messages": 0,
					"learning_plans": {},
				},
				"settings": {"language": "english", "theme": "light"},
			}
		)
		return True
	else:
		return False


def login_user(username: str, password: str) -> bool | str:
	"""
	Authenticates a user against the Firebase Realtime Database.

	This function checks if the provided username exists in the database and,
	if it does, verifies if the given password matches the stored password
	for that user.

	Args:
		username (str): The username to authenticate.
		password (str): The password to verify.

	Returns:
		bool|str:
		- True if the username and password match (successful login).
		- A string message (translated error) if the password is incorrect
		or the user is not found.
	"""
	current_users = users_ref.get() or {}

	if username in current_users:
		if current_users[username]["password"] == password:
			return True
		else:
			return translate_text("Wrong Password!", 0)
	else:
		return translate_text("User Not Found!", 1)


# Settings Functions
def change_name(username, new_name):
	users_ref.child(username).child("info").child("name").set(new_name)
	st.session_state.user_real_name = new_name
	return


def change_age(username, new_age):
	users_ref.child(username).child("info").child("age").set(new_age)
	st.session_state.age = new_age
	return


def change_username(current_username, new_username):
	users = users_ref.get()

	if current_username not in users:
		return False, translate_text("Current username not found!", 2)

	elif not check_username_validity(new_username):
		return (
			False,
			translate_text("Invalid username format. Only letters, numbers, underscores, and hyphens are allowed!", 3),
		)

	elif current_username == new_username:
		return

	elif new_username in users:
		return False, translate_text("Username already exists!", 4)

	else:
		# 1️⃣ Get all data of the current user
		user_data = users_ref.child(current_username).get()

		if user_data:
			users_ref.child(new_username).set(user_data)  # Set the new username data
			users_ref.child(current_username).delete()  # Delete the old username data
			st.session_state.username = (
				new_username  # Update the username session state
			)

			return True, translate_text("Username updated successfully!", 5)


def change_password(username, current_password, new_password, confirm_password):
	users = users_ref.get()
	user_current_password = users_ref.child(username).child("password").get()

	if not current_password or not new_password or not confirm_password:
		return False, translate_text("Please fill in all password fields", 6)

	elif username not in users:
		return False, translate_text("User not found!", 7)

	elif current_password != user_current_password:
		return False, translate_text("Current password is incorrect!", 8)

	elif current_password == new_password:
		return False, translate_text("New password can't be the same as the current password!", 9)

	elif not check_password_validity(new_password):
		return False, translate_text("Password must be at least 4 characters long!", 10)

	elif new_password != confirm_password:
		return False, translate_text("Passwords do not match!", 11)

	else:
		users_ref.child(username).child("password").set(new_password)
		return True, translate_text("Password updated successfully!", 12)


def delete_user(username: str):
	users = users_ref.get()
	if username in users:
		users_ref.child(username).delete()
		return True
	else:
		return False


def apply_theme(theme):
	config_toml_file_path = Path(".streamlit/config.toml")
	themes = default_data_ref.child("themes").get()

	# Load existing config
	if config_toml_file_path.exists():
		config = toml.load(config_toml_file_path)
	else:
		config = {}

	# Update only the theme section
	config["theme"] = themes[theme]

	# Save back to file
	config_toml_file_path.parent.mkdir(exist_ok=True)
	with open(config_toml_file_path, "w") as f:
		toml.dump(config, f)

	return


# Ask AI Functions
def add_chat_to_history(
	username: str,
	chat_num,
	role: str,
	message: str,
	attached_file_names: list = None,
	attached_image_files: list = None,
) -> None:
	
	"""
	Appends a new chat message (and optional attachments) to a user's AI chat history
	in the Firebase Realtime Database.

	The function stores messages under the user's `history/ai/chat{chat_num}/data`
	path. If the chat does not exist yet, it initializes it with an empty list.
	Attachments are stored alongside messages, with image files paired to their
	corresponding filenames.

	Args:
		username (str): The username of the user whose history is being updated.
		chat_num (str | int): The chat session number identifier (e.g., "1", "2").
		role (str): The role of the message sender (e.g., "user" or "model").
		message (str): The text content of the message.
		attached_file_names (list[str], optional): A list of filenames attached to the message.
			Files with extensions ["png", "jpeg", "webp"] are treated as images.
		attached_image_files (list[Any], optional): A list of image file objects corresponding
			to `attached_file_names`. The order must match the file list.

	Returns:
		None
	"""

	ai_chat_history_ref = (
		users_ref.child(username)
		.child("history")
		.child("ai")
		.child(f"chat{chat_num}")
		.child("data")
	)
	ai_history_list = ai_chat_history_ref.get()

	if ai_history_list is None:
		ai_history_list = []

	image_extensions_list = ["png", "jpeg", "webp"]
	images_index_count = 0

	if attached_file_names:
		attachments = []
		for file_name in attached_file_names:
			file_extension = file_name.split(".")[-1]

			if file_extension in image_extensions_list:
				attachments.append(
					{
						"file_name": file_name,
						"file": attached_image_files[images_index_count],
					}
				)
				images_index_count += 1
			else:
				attachments.append({"file_name": file_name})

		ai_history_list.append(
			{"role": role, "parts": [{"text": message}], "attachments": attachments}
		)
	else:
		ai_history_list.append({"role": role, "parts": [{"text": message}]})
		
	ai_chat_history_ref.set(ai_history_list)


def prompt_chat_history(username, chat_num, current_prompt: str = None):
	ai_history = (
		users_ref.child(username)
		.child("history")
		.child("ai")
		.child(f"chat{chat_num}")
		.child("data")
		.get()
	)

	if ai_history is None:
		ai_history = []

	ai_history_without_file_data = copy.deepcopy(ai_history)
	if ai_history_without_file_data is not None:
		for message in ai_history_without_file_data:
			if "attachments" in message:
				attachments_list = message["attachments"]
				for attachment_dict in attachments_list:
					if "file" in attachment_dict:
						del attachment_dict["file"]

	if current_prompt is None:
		current_prompt = ai_history[-1]["parts"][0]["text"]

	return_string = f"""
Previous Chat History In This Chat Session:
{ai_history_without_file_data}
Current Prompt: {current_prompt}
Note:
1. If the user asked for a summary give him a long one by default.
2. Continue the conversation normally without making any response to this sidetext. Don't even tell the user that you are instructed not to make any response to this sidetext.
"""
	return return_string


def reverse_index_to_original(reversed_index, list_length):
	return (list_length - 1) - reversed_index


def delete_chat(username, chat_number: str) -> None:
	"""
	Deletes a chat from Firebase and shifts all later chats down by one index.
	Example: delete chat 3 → chat 4 becomes chat 3, chat 5 becomes chat 4, etc.
	"""

	if chat_number == "all":
		users_ref.child(username).child("history").child("ai").set({})
		return

	chats_ref = users_ref.child(username).child("history").child("ai")
	chats = chats_ref.get() or {}

	# Step 1: Delete the selected chat
	chats_ref.child(f"chat{chat_number}").delete()

	# Step 2: Shift later chats down
	for i in range(int(chat_number) + 1, len(chats) + 1):
		old_chat_key = f"chat{i}"
		new_chat_key = f"chat{i-1}"
		old_chat_data = chats_ref.child(old_chat_key).get()

		if old_chat_data:
			chats_ref.child(new_chat_key).set(old_chat_data)
			chats_ref.child(old_chat_key).delete()
	
	return


def generate_chat_title(username, chat_num):
	first_message = (
		users_ref.child(username)
		.child("history")
		.child("ai")
		.child(f"chat{chat_num}")
		.child("data")
		.get()
	)

	if first_message:
		first_message = first_message[0]
	else:
		return

	prompt = f"""
	This is a history of a conversation:
	{first_message}
	You have to generate a short title for this conversation.
	Strict Notes: You must reply with the only the final answer (the title), no explaination or more words
	"""
	try:
		chat_title = st.session_stategemini_2_5_flash_lite.generate_content(
			prompt
		).text.strip()
	except:
		try:
			chat_title = st.session_state.gemini_2_0_flash.generate_content(
				prompt
			).text.strip()
		except:
			chat_title = "Chat"

	users_ref.child(st.session_state.username).child("history").child("ai").child(
		f"chat{chat_num}"
	).child("title").set(chat_title)

	return


def wrap_long_words(text: str, max_len=40):
	processed_lines = []
	# Process each line separately
	for line in text.splitlines():
		wrapped_parts = []
		# Split the line into words to check for long ones
		for word in line.split():
			# If a word is too long, manually wrap it
			if len(word) > max_len:
				for i in range(0, len(word), max_len):
					wrapped_parts.append(word[i : i + max_len])
			else:
				wrapped_parts.append(word)
		processed_lines.append(" ".join(wrapped_parts))
	return "\n".join(processed_lines)


def chats_history_checkup(username: str):
	"""
	Checks that the chats structure is correct and does not miss any attributes which can be caused by the
	user interupting the process. The function does the following:\n
	1. Checks if the number of messages in any chat is odd which means the user interrupted the ai while generating the last response, so it generates the last response.\n
	2. Checks if the chat title is missing in any chat and generates it.\n

	Args:
			username (str): The username of the user.
	"""

	chats_ref = users_ref.child(username).child("history").child("ai")
	chats = chats_ref.get()

	if chats:
		for chat in chats:
			current_chat_num = str(chat[4:])

			# Deleting any empty AI response
			if chats[chat]["data"][-1]["parts"][0]["text"] == "":
				users_ref.child(username).child("history").child("ai").child(
					chat
				).child("data")[-1].delete()

			# 1
			messages = chats[chat]["data"]
			if messages is not None:
				if len(messages) % 2 != 0:
					prompt = prompt_chat_history(username, current_chat_num)

					try:
						last_response = (
							st.session_stategemini_2_5_flash_lite.generate_content(
								prompt
							).text
						)
					except:
						try:
							last_response = (
								st.session_state.gemini_2_0_flash.generate_content(
									prompt
								).text
							)
						except:
							last_response = "Error"

					# Adding the AI's message to the history
					add_chat_to_history(
						username,
						chat_num=current_chat_num,
						role="model",
						message=last_response,
					)

			# 2
			title = chats[chat].get("title")
			if title is None:
				generate_chat_title(username, current_chat_num)

	return


if st.session_state.username and "ask ai" in st.session_state.current_page.lower():
	chats_history_checkup(st.session_state.username)


# Learning Plan and Dashboard Functions
def generate_learning_plan(topic: str, difficulty: str, days: str, age: str, plan_language: str, custom_prompt: str) -> str:
	"""Generate a comprehensive learning plan using AI"""
	
	prompt = f"""
	Create a detailed {days}-day learning plan for the topic: "{topic}"
	
	General:
	- Age: {age}
	- Difficulty Level: {difficulty}
	- Available Days: {days}
	- Plan Language: {plan_language}
	- AR Appropriate: true/false
	- AR Search Term: short, general search term for AR model or null

	For each day, provide:
	1. Daily Topic/Focus
	2. Reading Content (2-3 paragraphs of educational text)
	3. YouTube Search Query (for finding relevant videos)
	4. 5-7 Flashcard questions and answers
	5. 5 Quiz questions with multiple choice answers (3 at least if the user requested custom)
	6. Estimated time to complete (in minutes)
	
	Format your response as JSON with this exact structure:
	{{
		"plan_title": "Learning Plan Title",
		"total_days": {days},
		"difficulty": "{difficulty}",
		"ar": {{
			"can_be_represented_in_3d": true/false,
			"ar_search_term": "search term (In English, only one)/ None if can_be_represented_in_3d is false",
		}},
		"days": [
			{{
				"day": 1,
				"title": "Day Title",
				"content": "Educational content text...",
				"youtube_query": "clear search term for a youtube video in {plan_language}, search term 2" (You may make it 2 when the 2 videos may not be relevant. Mostly, it is one),
				"flashcards": [
					{{"question": "Q1?", "answer": "A1"}},
					{{"question": "Q2?", "answer": "A2"}}
				],
				"quiz": [
					{{
						"question": "Quiz question?",
						"options": ["A", "B", "C", "D"],
						"correct": 0
					}}
				],
				"estimated_time": int in minutes (expected small)
			}}
		]
	}}
	
	{f"This is a custom prompt from the user to follow: {custom_prompt}\n If this is not relevant or meaningless, ignore it." if custom_prompt else ""}

	Increase quizzes and choices difficulty **a little (so you are ceratin about your answers)** day by day.
	Make sure the content is age-appropriate and follows a logical progression.
	Only return the above structure, no more words or explanations.
	"""
	
	try:
		response = st.session_state.gemini_2_5_flash.generate_content(prompt)
		# Clean the response to extract JSON
		json_text = re.sub(r'```json|```', '', response.text).strip()

		return json.loads(json_text)
	except Exception as e:
		return None


def save_new_learning_plan(username: str, plan_data, include_ar: bool) -> None:
	"""Save learning plan to Firebase"""
	plans_dict = users_ref.child(username).child("history").child("learning_plans").get()
	new_plan_num = 1 if plans_dict is None else len(plans_dict) + 1

	users_ref.child(username).child("history").child("learning_plans").child(f"plan{new_plan_num}").set({
		"plan_data": plan_data,
		"current_day": 1,
		"start_date": str(datetime.date.today()),
		"completed_days": [],
		"daily_progress": {},
		"completed": False
	})

	if not include_ar:
		users_ref.child(username).child("history").child("learning_plans").child(f"plan{new_plan_num}").child("plan_data").child("ar").child("can_be_represented_in_3d").set(False)
		users_ref.child(username).child("history").child("learning_plans").child(f"plan{new_plan_num}").child("plan_data").child("ar").child("ar_search_term").set(None)
	return


def save_ar_data(username: str, ar_data, plan_num=None) -> None:
	plans_dict = users_ref.child(username).child("history").child("learning_plans").get()

	if plan_num is None:
		plan_num = len(plans_dict)

	users_ref.child(username).child("history").child("learning_plans").child(f"plan{plan_num}").child("plan_data").child("ar").child("ar_data").set(ar_data)
	return


def get_learning_plan(username: str, plan_num=None) -> dict:
	"""Get learning plan from Firebase"""
	plans_dict = users_ref.child(username).child("history").child("learning_plans").get()
	if plans_dict is None:
		return
	
	if plan_num is None:
		plan_num = len(plans_dict)

	return users_ref.child(username).child("history").child("learning_plans").child(f"plan{plan_num}").get()


def get_learning_plans_history(username: str) -> list:
	plans_titles = []
	plans_ref = users_ref.child(username).child("history").child("learning_plans")
	plans_dict = plans_ref.get()

	if plans_dict is None:
		return []

	for plan_key in plans_dict:
		plan = plans_ref.child(plan_key).get()
		title = plan["plan_data"]["plan_title"]
		plans_titles.append(title)

	return plans_titles


def delete_learning_plan(username: str, plan_num: str) -> None:
	plans_ref = users_ref.child(username).child("history").child("learning_plans")
	plans = plans_ref.get()

	# Step 1: Deleting the selected learning plan
	plans_ref.child(f"plan{plan_num}").delete()

	# Step 2: Shift later plans down
	for i in range(int(plan_num) + 1, len(plans) + 1):
		old_plan_key = f"plan{i}"
		new_plan_key = f"plan{i-1}"
		old_plan_data = plans_ref.child(old_plan_key).get()

		if old_plan_data:
			plans_ref.child(new_plan_key).set(old_plan_data)
			plans_ref.child(old_plan_key).delete()
	
	return


def update_daily_progress(username, day, section, completed=True) -> None:
	"""Update progress for a specific day and section"""

	plans_dict = users_ref.child(username).child("history").child("learning_plans").get()
	plan_num = 1 if plans_dict is None else len(plans_dict)

	users_ref.child(username).child("history").child("learning_plans").child(f"plan{plan_num}").child("daily_progress").child(f"day_{day}").child(section).set(completed)

	return


def update_current_day(username: str, completed_days_list: list, day: int, plan_num=None) -> None:
	"""Update current day in Firebase"""
	plans_dict = users_ref.child(username).child("history").child("learning_plans").get()
	
	plan_num = len(plans_dict) if not plan_num else plan_num

	users_ref.child(username).child("history").child("learning_plans").child(f"plan{plan_num}").child("completed_days").set(completed_days_list)
	users_ref.child(username).child("history").child("learning_plans").child(f"plan{plan_num}").child("current_day").set(day)

	return


def calculate_streak(username, plan_num=None):
	"""Calculate learning streak"""
	plan_data = get_learning_plan(username, plan_num)

	if not plan_data:
		return 0
	
	completed_days = plan_data.get("completed_days", [])
	if not completed_days:
		return 0
	
	# Sort completed days
	completed_days.sort()
	
	# Calculate consecutive days from the end
	streak = 0
	today = datetime.date.today()
	current_date = today
	
	for day in reversed(completed_days):
		completion_date = datetime.datetime.strptime(day, "%Y-%m-%d").date()
		if completion_date == current_date:
			streak += 1
			current_date -= datetime.timedelta(days=1)
		else:
			break
	
	return streak


def get_plan_badges(username: str, progress_percentage, plan_num=None):
	"""Get user badges based on progress"""
	badges = []
	plan_data = get_learning_plan(username, plan_num)
	
	if not plan_data:
		return badges
	
	completed_days = plan_data.get("current_day") - 1
	streak = calculate_streak(username, plan_num)
	
	# Day-based badges
	if completed_days >= 1:
		badges.append({"name": "First Steps", "icon": "🚀", "color": "#4CAF50"})
	if progress_percentage >= 25:
		badges.append({"name": "Quarter Way", "icon": "🎯", "color": "#00BCD4"})
	if progress_percentage >= 50:
		badges.append({"name": "Halfway Hero", "icon": "⚡", "color": "#FF9800"})
	if progress_percentage >= 100:
		badges.append({"name": "Plan Master", "icon": "👑", "color": "#9C27B0"})
	
	# Streak-based badges
	if streak >= 2:
		badges.append({"name": "Streak Starter", "icon": "🔥", "color": "#F44336"})
	if streak >= 7:
		badges.append({"name": "Week Streaker", "icon": "⚡", "color": "#FF5722"})
	if streak >= 14:
		badges.append({"name": "Habit Builder", "icon": "💪", "color": "#3F51B5"})
	
	return badges


def get_motivational_quote():
	"""Get a random motivational quote"""

	quotes = [
		"The expert in anything was once a beginner. - Helen Hayes",
		"Learning never exhausts the mind. - Leonardo da Vinci",
		"Education is the passport to the future. - Malcolm X",
		"The more that you read, the more things you will know. - Dr. Seuss",
		"Live as if you were to die tomorrow. Learn as if you were to live forever. - Mahatma Gandhi",
		"An investment in knowledge pays the best interest. - Benjamin Franklin",
		"The capacity to learn is a gift; the ability to learn is a skill. - Brian Herbert",
		"Learning is a treasure that will follow its owner everywhere. - Chinese Proverb"
	]
	return random.choice(quotes)


# AR, Content and Quiz Generation Functions
def generate_ar(topic_name: str) -> list | str:
	# Step 1: Declaring API keys variables, getting things ready

	# Sketchfab
	SKETCHFAB_API_KEY = st.secrets["SKETCHFAB_TOKEN"]
	AUTHORIZATION_HEADER = {"Authorization": f"Token {SKETCHFAB_API_KEY}"}

	# Echo3D
	ECHO3D_API_KEY = st.secrets["ECHO3D_API_KEY"]
	ECHO3D_SECURITY_KEY = st.secrets["ECHO3D_SECURITY_KEY"]
	ECHO3D_USER_AUTH_KEY = st.secrets["ECHO3D_USER_AUTHENTICATION_KEY"]
	ECHO3D_EMAIL = st.secrets["ECHO3D_EMAIL"]

	# Filtering the topic name to avoid errors
	topic_name = re.sub(r'[\\/*?:"<>|]', "", topic_name).lower()

	# Step 2: Searching for 3D models from Sketchfab, a huge 3D Library
	def search_models(query, quantity=20):
		sketchfab_search_endpoint = "https://api.sketchfab.com/v3/search"
		params = {
			"q": query,
			"type": "models",
			"downloadable": True,
			"count": quantity,
		}

		response = requests.get(
			sketchfab_search_endpoint, headers=AUTHORIZATION_HEADER, params=params
		)
		response.raise_for_status()
		return response.json()["results"]

	try:
		models = search_models(topic_name)
	except Exception as e:
		return False, f"An Error Occured: {e}"
	
	if not models:
		return False, "No models found for this topic."

	# Step 3: Detrmining the best describing image by Gemini

	# Determining the best image function
	def best_model(images: list, topic_name: str) -> int:
		"""
		Analyzes multiple images with the Gemini API, allowing you to name them.

		Args:
			images: A list of dictionaries, where each dict has 'bytes', 'mime_type', and 'name'.
			text_prompt: The text to accompany the images.

		Returns:
			int:
			The index of the best model from the models list.
		"""

		contents = []
		
		images_names = [item['name'] for item in images]
		prompt_with_names = f"""
Analyis these images, which of these images is the best describing/showing a {topic_name}.
The evaluation terms are:
1. The realeativeness of the image to a {topic_name}, content, structure, color, etc.
2. The quality of the image.
The Images Names: {", ".join(images_names)} (Note: They are in the same order as the provided images)
Strict Notes:
1. The return answer must be the number of the chosen image only, no more words or explainations.
2. If no image matches or the images are all really bad, the return answer must be "None"
3. You must return an answer
"""
		# Adding each image's data to the contents
		for image_data in images:
			image_part = {
				"mime_type": image_data["mime_type"],
				"data": image_data["bytes"]
			}
			contents.append(image_part)

		# Adding the prompt to the contents after the images
		contents.append(prompt_with_names)

		# Generating the response
		response = st.session_state.gemini_2_0_flash.generate_content(contents)
		try:
			model_index = int(response.text) - 1
		except:
			# Getting the first model by default if an error occurs
			model_index = 0

		return model_index
	

	# Getting the thumbnails images bytes, mime types
	thumbnails_images_urls = [model["thumbnails"]["images"][0]["url"] for model in models]
	thumbnails_images_bytes = []
	thumbnails_images_mime_types = []

	for i, image_url in enumerate(thumbnails_images_urls):
		try:
			# Getting the image bytes
			response = requests.get(image_url)
			thumbnails_images_bytes.append(response.content)

			# Getting the image mime type
			mime_type, _ = mimetypes.guess_type(image_url)
			thumbnails_images_mime_types.append(mime_type)

		except Exception as e:
			pass

	# Creating the image data
	images_data = [
		{"bytes": image_bytes, "mime_type": mime_type, "name": f"image_{i+1}"}
		for i, (image_bytes, mime_type) in enumerate(zip(thumbnails_images_bytes, thumbnails_images_mime_types))
	]

	# Determining the best model
	best_model_try = best_model(images_data, topic_name)

	if best_model_try is None:
		best_model_index = 0  # OR 3D models generation code here
	else:
		best_model_index = best_model_try

	# Step 4: Downloading the model
	model_uid = models[best_model_index]["uid"]
	model_embed_url = models[best_model_index]["embedUrl"]
	sketchfab_download_endpoint = f"https://api.sketchfab.com/v3/models/{model_uid}/download"

	# Sending the request to get the downloads options
	response = requests.get(sketchfab_download_endpoint, headers=AUTHORIZATION_HEADER)

	if response.status_code == 200:
		# Downloading the model's GLB
		model_glb_url = response.json().get("glb", {}).get("url")

		if model_glb_url:
			try:
				response = requests.get(model_glb_url)
				response.raise_for_status()
				model_glb_bytes = response.content
			except Exception as e:
				return False, f"An Error Occured: {e}"
			
		else:
			return False, "An Error Occured: No GLB file found"
		
	else:
		return False, f"An Error Occured: {response.status_code}"
	
	# Step 5: Ceating the AR experience on Echo3D with the 3D model

	# Setting up the request payload
	echo3d_uploaded_endpoint = "https://api.echo3D.com/upload"
	files = {"file_model": (f"{topic_name}.glb", model_glb_bytes)}
	data = {
		"key": ECHO3D_API_KEY,
		"email": ECHO3D_EMAIL,
		"userKey": ECHO3D_USER_AUTH_KEY,
		"target_type": 2,  # Have 0 in a quick video to show the judges
		"hologram_type": 2,
		"secKey": ECHO3D_SECURITY_KEY,
		"type": "upload",
	}

	# Sending the request to upload the model to Echo3D
	try:
		response = requests.post(echo3d_uploaded_endpoint, files=files, data=data)
		response.raise_for_status()

		if response.status_code != 200:
			return False, f"An Error Occured: {response.status_code}"

	except Exception as e:
		traceback.print_exc()
		return False, f"An Error Occured: {e}"
	
	# Getting the link of the AR experience
	ar_experience_url = response.json().get("additionalData", {}).get("shortURL")

	if not ar_experience_url:
		return False, "An Error Occured: No AR experience URL found"

	# Step 6: Returning the AR link and QR code
	return True, [ar_experience_url, model_embed_url]


def get_video_id(queries: list) -> list:
	SERPAPI_API_KEY = st.secrets["SERPAPI_API_KEY"]
	ids = []

	for query in queries:			
		params = {
		"engine": "youtube",
		"search_query": query,
		"api_key": SERPAPI_API_KEY
		}

		search = GoogleSearch(params)
		results = search.get_dict()
		video_results = results["video_results"]

		try:
			video_url = video_results[0]["link"]

			video_id = video_url.split("v=")[1]
			ids.append(video_id)
		except:
			return None
	
	return ids


def generate_content(
	topic_name: str,
	age: str,
	grade: str=None,
	language: str="English",
	length: str="auto",
	complexity: str="auto",
	custom_prompt: str=None,
	model_embed_url: str=None
) -> list | str:

	# Step 1: Declaring API keys variables, getting things ready
	SERPAPI_API_KEY = st.secrets["SERPAPI_API_KEY"]

	HASHNODE_ACCESS_TOKEN = st.secrets["HASHNODE_ACCESS_TOKEN"]
	PUBLICATION_HOST = "eureka1.hashnode.dev"
	HASHNODE_API_URL = "https://gql.hashnode.com/"

	# Step 2: Preparing for the prompt
	
	# Getting relevant images
	prompt = f"""
If you were to make an article about {topic_name}, make a discription of two images you may use (comma-separated). Max: 2
Example output if the topic is Human Heart: Human Heart, Human Heart Chambers
"""
	
	response = st.session_state.gemini_2_0_flash.generate_content(prompt)
	image_descriptions = response.text

	queries = [description.strip() for description in image_descriptions.split(",")]

	# Images from Google Images by SerpAPI
	def get_images(queries: list) -> str:
		urls = {}

		for i, query in enumerate(queries):
			params = {
			"engine": "google_images",
			"q": query,
			"api_key": SERPAPI_API_KEY
			}

			search = GoogleSearch(params)
			results = search.get_dict()

			url_index = i
			url = results["images_results"][url_index]["original"]
			if url in urls:
				url_index += 1
				url = results["images_results"][url_index]["original"]

			urls[query] = url

		return urls
		
	images_urls = get_images(queries)

	# Step 3: Generating the content using Gemini and getting the article data for Hashnode
	telegraph_prompt = f"""Generate an article (content) about {topic_name} suitable for age {age} {f"and grade: {grade}" if grade else ''} in terms of words complexity, content length, etc.
It must be written in the Markdown Style.
Content Language: {language.capitalize()}
{f"Content Length: {length.title()}" if length != "auto" else ""}
{f"Content Complexity: {complexity.title()}" if complexity != "auto" else ""}
{f"Custom Prompt From The User: {custom_prompt}. If it is not relevant or meaningful, ignore it and don't mention it." if custom_prompt else ""}
At the end, it is preferable to add a 'Did you know' section as well as a 'What is next?' section, including 'If you learned about ...., so you might also like: .....' with real links to other articles.
{f"Here are images to use: {images_urls} by '![Alt Text](Image URL)'" if images_urls else ""}
Output structure:
Title: ...
Subtitle: ...
Youtube Video Search Query: ... (The exact as the given article topic, but cleared from typos)
(Markdown formatted text here, you can make new lines in it)

1. Only return the output article in that structure, no more words or explainations.
2. The markdow data should start at line 3.
"""

	response = st.session_state.gemini_2_0_flash.generate_content(telegraph_prompt).text

	# Extracting the data
	# Setting default values
	title = topic_name
	subtitle = None
	youtube_video_query = None
	
	content_starting_index = 4
	content = ""

	# Extracting the title, subtitle and content
	lines = response.splitlines()
	for i, line in enumerate(lines):
		if line.startswith("Title"):
			title = line.split(':', 1)[1].strip()
		elif line.startswith("Subtitle"):
			subtitle = line.split(':', 1)[1].strip()
		elif line.lower().startswith("youtube video search query"):
			youtube_video_query = line.split(':', 1)[1].strip()
			content_starting_index = i + 1
	
	# 3D Model Embed
	if model_embed_url:
		content += f"""<iframe title="Sketchfab Model" frameborder="0" allowfullscreen 
mozallowfullscreen="true" webkitallowfullscreen="true" style="width: 100%; height: 480px;" 
src="{model_embed_url}"></iframe>

"""
	content += "\n".join(lines[content_starting_index:])
	content_for_doc = "\n".join(lines[content_starting_index:])

	# Youtube video embed
	if not youtube_video_query:
		youtube_video_query = topic_name

	video_id = get_video_id([youtube_video_query])[0]

	if video_id:
		content += f"""

### Don't just read, experience it! Check out this video:
<iframe src="https://www.youtube.com/embed/{video_id}" frameborder="0" allowfullscreen style="width: 100%; height: 480px;"></iframe>
"""

	# Step 4: Creating the article on Hashnode
	def create_hashnode_article(title: str, subtitile: str, content_markdown: str, publication_host: str, access_token: str):
		"""
		Sends a GraphQL request to create a new article on Hashnode.

		Args:
			title (str): The title of the article.
			subtitle (str): The subtitle of the article.
			content_markdown (str): The article content in Markdown format.
			publication_host (str): The host/domain of the publication.
			access_token (str): Your Hashnode Personal Access Token.

		Returns:
			dict: The JSON response from the Hashnode API, or None if an error occurred.
		"""
		
		# Set up the request headers for authentication
		headers = {
			"Authorization": access_token,
			"Content-Type": "application/json",
		}

		# GraphQL mutation to create a new post
		graphql_query = """
		mutation PublishPost($input: PublishPostInput!) {
		publishPost(input: $input) {
			post {
			id
			slug
			url
			publication {
				id
				url
			}
			}
		}
		}
		"""

		# Variables (payload) for the mutation to make the new post
		variables = {
			"input": {
				"title": title,
				"contentMarkdown": content_markdown,
				"publicationId": None,
				"tags": [],
				"seriesId": None,
				"coverImageOptions": None,
				"metaTags": None,
				"subtitle": subtitile,
			}
		}

		# Getting the publication ID using the host
		publication_id_query = f"""
		query Publication {{
		publication(host: "{publication_host}") {{
			id
		}}
		}}
		"""

		# Send the request to get the publication ID
		try:
			response_pub_id = requests.post(
				HASHNODE_API_URL, headers=headers, json={"query": publication_id_query}
			)
			response_pub_id.raise_for_status() # Raise an HTTPError for bad responses (4xx or 5xx)
			pub_id_data = response_pub_id.json()
			
			if "errors" in pub_id_data:
				return False
			
			publication_id = pub_id_data["data"]["publication"]["id"]
			variables["input"]["publicationId"] = publication_id

		except Exception as e:
			return False

		# Sending the request to publish the article
		try:
			response = requests.post(
				HASHNODE_API_URL, headers=headers, json={"query": graphql_query, "variables": variables}
			)
			response.raise_for_status() # Raise an HTTPError for bad responses (4xx or 5xx)
			return response.json()
		except Exception as e:
			return False
	
	# Getting the page URL
	page_data = create_hashnode_article(title, subtitle, content, PUBLICATION_HOST, HASHNODE_ACCESS_TOKEN)
	page_url = page_data["data"]["publishPost"]["post"]["url"]

	# Step 5: Creating the QR code for the page

	# Creating the QR code
	qr_code = qrcode.make(page_url)
	# Saving the QR code to a buffer
	qr_code_buffer = BytesIO()
	qr_code.save(qr_code_buffer, "PNG")
	# Rewend the buffer to the beginning
	qr_code_buffer.seek(0)

	# Step 6: Creating the document
	def strip_markdown(md_text):
		# Convert markdown → HTML
		html = markdown.markdown(md_text)

		# Use BeautifulSoup to get plain text
		soup = BeautifulSoup(html, "html.parser")
		text_content = soup.get_text()

		return text_content

	
	content_passage = strip_markdown(content_for_doc)

	# Create the Word document
	content_doc = Document()

	# Add a centered bold title
	content_doc_heading = content_doc.add_heading(level=1)
	content_doc_run = content_doc_heading.add_run(f"{topic_name.title()} for Age {age}")
	content_doc_run.font.name = "Arial"
	content_doc_run.font.size = Pt(20)
	content_doc_run.bold = True
	content_doc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

	# Add spacing
	content_doc.add_paragraph("")

	# Add the full cleaned passage, paragraph by paragraph
	for line in content_passage.strip().split("\n"):
		if line.strip():  # Skip empty lines
			content_doc_para = content_doc.add_paragraph()
			content_doc_run = content_doc_para.add_run(line.strip())
			content_doc_run.font.name = "Arial"
			content_doc_run.font.size = Pt(16)
			if is_rtl(language):
				content_doc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
			else:
				content_doc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

	# Save to memory
	content_doc_buffer = BytesIO()
	content_doc.save(content_doc_buffer)
	content_doc_buffer.seek(0)

	# Step 7: Returning the results
	return [page_url, qr_code_buffer, content_doc_buffer]


def generate_quiz(
	topic_name: str,
	questions_num: int,
	difficulty: str,
	langauge: str,
	questions_types: list=None,
):
	# Step 1: Generating the questions, choices and correct answers by Gemini
	quiz_prompt = f"""
Create a quiz of {questions_num} questions about {topic_name}.
Language: {langauge}, Difficulty: {difficulty}.
All questions types are multiple choice / multiple select / true or false / fill in the blank / essay. In this quiz you will {'only use' + ' / '.join(questions_types) if questions_types else 'use all of them'}.
You must return the quiz data in this form:
Quiz title: ... (Don't write 'Quiz' or 'Form' at the end)
Question type: ...
Question: ......
Choices: (Only for multiple choice, multiple select and true or false questions)
Choice A: ...
Choice B: ...
Choice C: ...
Choice D: ...
Correct Answer: ... (e.g. C / A, B, D for multiple choice, multiple select and true or false questions | The answer essay and fill in the blank questions)
Explaination If Answered Wrong: ... (A short, rich explanation if the answer is wrong)

{'Note: Make a good ratio for using all the question types.' if questions_types and len(questions_types) > 1 else ''}
Note:
1. For fill in the blank questions, you can only add one blank.
2. For essay questions, make the answer dynamic as you like. 
3. You must return only the quiz data in that structure, no more words or explainations.
"""
	{'question_type': 'fill in the blank', }


# Step 3: Creating the login/signup page
if not st.session_state.logged_in:
	st.session_state.current_page = "🧠 Ask AI"

	st.title("🔐 " + translate_text("Welcome to Eureka, your full educational station!", 13))

	if st.session_state.login_or_signup == "login":
		with st.container(border=True, key="login_container"):
			with st.form(border=False, key="login_form"):
				st.subheader(translate_text("Login", 14))

				username_input = st.text_input(translate_text("Username", 15), key="login_user")
				password_input = st.text_input(
					translate_text("Password", 16), type="password", key="login_pass", autocomplete="off"
				)

				if st.form_submit_button(
					translate_text("Login", 17), use_container_width=True, type="primary"
				):
					if username_input and password_input:
						login_try = login_user(username_input, password_input)
						if login_try == True:
							st.success("✅ " + translate_text("Login successful!", 18))
							st.session_state.logged_in = True
							st.session_state.username = username_input
							st.session_state.user_real_name = (
								users_ref.child(st.session_state.username)
								.child("info")
								.child("name")
								.get()
							)
							st.session_state.age = (
								users_ref.child(st.session_state.username)
								.child("info")
								.child("age")
								.get()
							)
							st.session_state.theme = (
								users_ref.child(st.session_state.username)
								.child("settings")
								.child("theme")
								.get()
							)
							apply_theme(st.session_state.theme)
							st.rerun()
						else:
							st.error(f"❌ {login_try}")
					else:
						st.error("❌ " + translate_text("Please fill in all the fields.", 19))

			" "
			if st.button(translate_text("Don't have an account? **Sign Up**!", 20), type="tertiary"):
				st.session_state.login_or_signup = "signup"
				st.rerun()

	elif st.session_state.login_or_signup == "signup":
		with st.container(border=True):
			with st.form(border=False, key="signup_form"):
				st.subheader(translate_text("Sign Up", 27))

				col1, col2 = st.columns(2)
				with col1:
					first_name_input = st.text_input(translate_text("First Name", 92), key="signup_first_name")
				with col2:
					last_name_input = st.text_input(translate_text("Last Name", 93), key="signup_last_name")

				col1, col2 = st.columns([0.75, 0.25])
				with col1:
					new_username_input = st.text_input(
						translate_text("Username", 23),
						key="signup_user",
						placeholder=translate_text("Only letters, numbers, underscores and hyphens are allowed.", 24),
					)
				with col2:
					new_age_input = st.number_input(
						translate_text("Age", 22), min_value=3, max_value=120, key="signup_age"
					)

				new_password_input = st.text_input(
					translate_text("Create a Password", 25),
					type="password",
					key="signup_pass",
					autocomplete="off",
				)
				st.write(translate_text("- Password must be at least 4 characters long.", 26))

				if st.form_submit_button(
					translate_text("Sign Up", 27), use_container_width=True, type="primary"
				):

					if (
						first_name_input
						and last_name_input
						and new_age_input
						and new_username_input
						and new_password_input
					):
						# Making up the full name
						new_name = first_name_input + " " + last_name_input

						if check_username_validity(new_username_input):

							if check_password_validity(new_password_input):

								if register_user(
									new_username_input,
									new_password_input,
									first_name_input + " " + last_name_input,
									new_age_input,
								):
									st.success("✅ " + translate_text("Registered successfully!", 28))
									st.session_state.logged_in = True
									st.session_state.username = new_username_input
									st.session_state.age = new_age_input
									st.session_state.user_real_name = new_name
									st.rerun()

								else:
									st.error("❌ " + translate_text("Username already exists.", 29))

							else:
								st.error("❌ " + translate_text("Password is not strong enough.", 30))

						else:
							st.error(
								"❌ " + translate_text("Invalid username. Only letters, numbers, underscores and hyphens are allowed.", 31)
							)

					else:
						st.error("❌ " + translate_text("Please fill in all the fields.", 32))

			" "
			if st.button(translate_text("Already have an account? **Login**!", 33), type="tertiary"):
				st.session_state.login_or_signup = "login"
				st.rerun()

	st.stop()

# -----------------------------
# Sidebar Navigation
# -----------------------------

# Settings Icon Button
with st.sidebar:
	col1, col2 = st.columns(2)

	with col1:
		if st.button(translate_text("Home", 34), icon="🏠", use_container_width=True, key="sidebar_home_btn"):
			st.session_state.current_page = "🧠 Ask AI"
			st.session_state.new_chat = True
			st.rerun()

	with col2:
		if st.button(translate_text("Settings", 35), icon="⚙️", use_container_width=True, key="sidebar_settings_btn"):
			st.session_state.current_page = "⚙️ Settings"
			st.rerun()


	# Selecting the initial page
	if st.session_state.current_page != "⚙️ Settings":
		try:
			name = st.session_state.user_real_name.split(" ")[0]
		except:
			name = st.session_state.user_real_name

		st.title(f"👋 " + f"{translate_text("Welcome", 36)}, :red[{google_translate(name, st.session_state.language).title()}]")

		features_list_english = [
			"🧠 Ask AI",
			"📚 Learning Plans",
			"🏆 Learning Dashboard",
			"🔖 AR, Content and Quiz Package",
			"📄 Upload PDF & Summarize",
			"❓ Quiz Generator",
		]
		features_list = [
			f"🧠 " + translate_text("Ask AI", 38),
			f"📚 " + translate_text("Learning Plans", 86),
			f"🏆 " + translate_text("Learning Dashboard", 87),
			f"🔖 " + translate_text("AR, Content and Quiz Package", 39),
			f"📄 " + translate_text("Upload PDF & Summarize", 40),
			f"❓ " + translate_text("Quiz Generator", 41),
		]

		current_page = st.selectbox(
			translate_text("Choose a Feature", 37),
			features_list,
			index=features_list_english.index(st.session_state.current_page),
		)

		page_index = features_list.index(current_page)
		st.session_state.current_page = features_list_english[page_index]

	# Learning Plans History
	if "current_plan_num" not in st.session_state:
		plans_dict = users_ref.child(st.session_state.username).child("history").child("learning_plans").get()
		st.session_state.current_plan_num = len(plans_dict) if plans_dict else None
		
	if st.session_state.current_page == "📚 Learning Plans" or st.session_state.current_page == "🏆 Learning Dashboard":
		plans_titles = list(reversed(get_learning_plans_history(st.session_state.username)))

		if plans_titles:
			st.write("Learning Plans")

			col1, col2 = st.columns([0.85, 0.15])

			for i, plan_title in enumerate(plans_titles):
				current_plan_num = reverse_index_to_original(i, len(plans_titles)) + 1

				plan_completed = users_ref.child(st.session_state.username).child("history").child("learning_plans").child(f"plan{current_plan_num}").child("completed").get()
				print(f"plan_completed: {plan_completed} for {plan_title}")
				with col1:
					with stylable_container(
						key=f"sidebar_chat_button_{i}",
						css_styles="""
						button{
							display: flex;
							justify-content: flex-start;
							width: 100%;
						}
						""",
					):
						
						if st.button(plan_title, icon="✅" if plan_completed else None, key=f"plan_btn_{i}", use_container_width=True):
							st.session_state.current_page = "🏆 Learning Dashboard"
							st.session_state.current_plan_num = current_plan_num
							st.session_state.start_learning = True
							st.rerun()
						
						" "
				
				with col2:
					if st.button("🗑️", key=f"delete_plan_btn_{i}", use_container_width=True):
						delete_learning_plan(st.session_state.username, current_plan_num)
						st.rerun()
					
					" "

# Fixing the sidebar width
sidebar_width = 300

st.markdown(
    f"""
    <style>
        section[data-testid="stSidebar"] {{
            width: {sidebar_width}px !important;
        }}
    </style>
    """,
    unsafe_allow_html=True,
)

# -----------------------------
# ⚙️ Settings
# -----------------------------

if st.session_state.current_page == "⚙️ Settings":
	st.title("⚙️ " + translate_text("Settings", 42))

	# Create tabs for different settings categories
	tab1, tab2 = st.tabs(["🎨 " + translate_text("Preferences", 43), "👤 " + translate_text("Account", 44)])

	# Tab 1: Preferences Settings
	with tab1:
		# Theme Settings
		st.subheader("🎨 " + translate_text("Theme Settings", 45))

		themes = default_data_ref.child("themes").get()
		st.session_state.theme = (
			users_ref.child(st.session_state.username)
			.child("settings")
			.child("theme")
			.get()
		)

		# Theme selection
		selected_theme = st.selectbox(
			translate_text("Theme", 46),
			[theme.capitalize() for theme in list(themes.keys())],
			index=list(themes.keys()).index(st.session_state.theme.lower()),
		)

		st.divider()

		# Language Settings
		st.subheader("🌐 " + translate_text("Language Settings", 84))

		# Language selection
		languages = [
			"English",
			"Spanish",
			"French",
			"German",
			"Italian",
			"Portuguese",
			"Russian",
			"Chinese",
			"Japanese",
			"Korean",
			"العربية",
		]
		current_lang = st.session_state.language
		if current_lang.lower() == "arabic":
			current_lang = "العربية"

		starting_index = languages.index(current_lang.capitalize())
		selected_language = st.selectbox(
			translate_text("Language", 47),
			languages,
			index=starting_index
		)

		if selected_language == "العربية":
			selected_language = "Arabic"

		st.divider()

		if st.button("💾 " + translate_text("Save", 48), key="language_save_btn"):
			# Theme
			st.session_state.theme = (
				selected_theme.lower()
			)  # Update the theme session state

			users_ref.child(st.session_state.username).child("settings").child(
				"theme"
			).set(
				st.session_state.theme
			)  # Save selected theme in the database

			apply_theme(st.session_state.theme)

			# Language
			st.session_state.language = (
				selected_language.lower()
			)  # Update the language session state

			users_ref.child(st.session_state.username).child("settings").child(
				"language"
			).set(
				st.session_state.language
			)  # Save selected language in the database
			st.rerun()

	# Tab 2: Account Settings
	with tab2:
		# Change Name Section
		st.subheader("📄 " + translate_text("Change Name", 49))

		# Current name
		st.info(f"{translate_text("Current Name", 85)}: {st.session_state.user_real_name}")

		with st.expander(translate_text("Change Name", 50)):
			with st.form(key="change_name_form", border=False):
				new_name = st.text_input(translate_text("New Name:", 51), key="new_name_input")

				if st.form_submit_button("📄 " + translate_text("Change Name", 52)) and new_name:
					change_name(st.session_state.username, new_name.strip())
					st.success("✅ " + translate_text("Name changed successfully!", 53))
					st.rerun()

		st.divider()

		# Change Age Section
		st.subheader("📄 " + translate_text("Change Age", 88))

		# Current age
		st.info(f"{translate_text("Current Age", 89)}: {st.session_state.age}")

		with st.expander(translate_text("Change Age", 88)):
			with st.form(key="change_age_form", border=False):
				new_age = st.number_input(translate_text("New Age", 90), key="new_age_input", min_value=3, max_value=120, step=1, value=int(st.session_state.age))

				if st.form_submit_button("📄 " + translate_text("Change Age", 88)) and new_age:
					change_age(st.session_state.username, new_age)
					st.success("✅ " + translate_text("Age changed successfully!", 91))
					st.rerun()
		
		st.divider()

		# Change Username Section
		st.subheader("🔄 " + translate_text("Change Username", 54))

		# Current username display
		st.info(f"{translate_text("Current Username", 55)}: {st.session_state.username}")

		with st.expander(translate_text("Change Username", 56)):
			with st.form(key="change_username_form", border=False):
				new_username = st.text_input(f"{translate_text("New Username", 57)}:", key="new_username_input")

				if (
					st.form_submit_button("🔄 " + translate_text("Change Username", 58))
					and new_username
				):
					success, return_message = change_username(
						st.session_state.username, new_username.strip()
					)
					if success:
						st.success(f"✅ {return_message}")
						st.rerun()
					elif success is None:
						pass
					else:
						st.error(f"❌ {return_message}")

		st.divider()

		# Change Password Section
		st.subheader("🔒 " + translate_text("Change Password", 59))

		with st.expander(translate_text("Change Password", 60)):
			current_password = st.text_input(
				f"{translate_text("Current Password", 61)}:", type="password", key="old_password_input", value=""
			)

			new_password = st.text_input(
				f"{translate_text("New Password", 62)}:", type="password", key="new_password_input", value=""
			)

			confirm_password = st.text_input(
				f"{translate_text("Confirm New Password", 63)}:",
				type="password",
				key="confirm_password_input",
				value="",
			)

			if st.button("🔒 " + translate_text("Change Password", 64), key="change_password_btn"):
				success, return_message = change_password(
					st.session_state.username,
					current_password,
					new_password,
					confirm_password,
				)
				if success:
					st.success(f"✅ {return_message}")
					st.rerun()
				else:
					st.error(f"❌ {return_message}")

		st.divider()

		# Sign Out/Delete Account Section
		st.subheader(translate_text("Actions", 65))

		col1, col2 = st.columns([1, 1])

		with col1:
			if st.button(
				translate_text("Sign Out", 66),
				key="sign_out_btn",
				use_container_width=True,
				icon="🚪",
			):
				st.session_state.logged_in = False
				st.session_state.username = None
				st.session_state.user_real_name = None

				st.success("✅ " + translate_text("Successfully signed out!", 67))
				st.rerun()

		with col2:
			@st.dialog("Are you sure you want to delete your account?")
			def confirm_delete():
				st.write("**Your account will be permanently deleted.**")
				cd_col1, cd_col2 = st.columns(2)

				with cd_col1:
					if st.button("Yes", type="secondary", use_container_width=True):
						deleted = delete_user(st.session_state.username)
						if deleted:
							st.session_state.logged_in = False
							st.session_state.login_or_signup = "login"
							st.session_state.username = None
							st.session_state.user_real_name = None
							st.rerun()
						else:
							st.error(
								"❌ " + translate_text("An error occurred while deleting the account.", 72)
							)

				with cd_col2:
					if st.button("No", type="primary", use_container_width=True):
						st.rerun()

							
			if st.button(
				translate_text("Delete Account", 68),
				key="delete_account_btn",
				use_container_width=True,
				icon="🗑️",
			):
				confirm_delete()

		st.divider()

		if st.button(
			translate_text("Restore All Settings", 74),
			type="primary",
			key="restore_all_settings_btn",
			use_container_width=True,
			icon="🔄"
		):
			users_ref.child(st.session_state.username).child("settings").set(
				{"language": "english", "theme": "light"}
			)
			st.session_state.language = "english"
			st.session_state.theme = "light"
			apply_theme(st.session_state.theme)

			st.success("✅ " + translate_text("All settings restored to default!", 75))
			st.rerun()

		st.divider()

		# Account Statistics
		st.subheader("📊 " + translate_text("Account Statistics", 76))

		col1, col2, col3 = st.columns(3)

		total_messages = 0
		chats = (
			users_ref.child(st.session_state.username)
			.child("history")
			.child("ai")
			.get()
		)

		if chats:
			for chat in chats:
				chat_data_list = chats[chat]["data"]
				for msg_dict in chat_data_list:
					if msg_dict["role"] == "user":
						total_messages += 1

		with col1:
			st.metric(translate_text("Total Messages", 77), total_messages)

# -----------------------------
# 🧠 AI Assistant
# -----------------------------

elif "ask ai" in st.session_state.current_page.lower():
	# New Chat Button
	if st.sidebar.button(
		translate_text("New Chat", 78),
		icon="✍️",
		key="new_chat_btn",
		use_container_width=True,
		type="secondary",
	):
		st.session_state.new_chat = True
		st.rerun()

	# Displaying the previous chats history in the sidebar
	chats_titles = []
	all_ai_history = (
		users_ref.child(st.session_state.username).child("history").child("ai").get()
	)

	if all_ai_history:
		st.sidebar.write(translate_text("Recent", 79))

		for chat in all_ai_history:
			if "title" in all_ai_history[chat]:
				title = all_ai_history[chat]["title"]
				chats_titles.append(title)

		chats_titles_len = len(chats_titles)
		chats_titles_reversed = list(reversed(chats_titles))  # Newest First

		for i, title in enumerate(chats_titles_reversed):
			col1, col2 = st.sidebar.columns([0.85, 0.15])
			with col1:
				with stylable_container(
					key=f"sidebar_chat_button_{i}",
					css_styles="""
					button{
						display: flex;
						justify-content: flex-start;
						width: 100%;
					}
					""",
				):
					if st.button(title, key=f"chat_btn_{i}", use_container_width=True):
						# Getting that chat's number
						st.session_state.current_chat_num = str(
							reverse_index_to_original(i, chats_titles_len) + 1
						)
						st.session_state.new_chat = False
						st.session_state.chat_history = (
							users_ref.child(st.session_state.username)
							.child("history")
							.child("ai")
							.child(f"chat{st.session_state.current_chat_num}")
							.child("data")
							.get()
							or []
						)
						st.rerun()
					" "
			with col2:
				if st.button("🗑️", key=f"delete_chat_btn_{i}", use_container_width=True):
					chat_num_to_delete = str(
						reverse_index_to_original(i, chats_titles_len) + 1
					)
					delete_chat(st.session_state.username, chat_num_to_delete)
					st.session_state.new_chat = True
					st.rerun()

	# Displaying the chat history
	docs_extensions_list = ["pdf", "txt"]
	image_extensions_list = ["png", "jpeg", "webp"]
	video_extensions_list = ["mp4", "webm", "flv", "mov", "mpg", "wmv", "3gpp", "mpeg"]
	audio_extensions_list = ["aac", "flac", "mp3", "m4a", "mpga", "opus", "pcm", "wav"]

	if not st.session_state.new_chat:
		for msg_dict in st.session_state.chat_history:
			if msg_dict["role"] == "user":
				attachments = msg_dict.get("attachments")

				if attachments:
					col1, col2 = st.columns([3, 1])
					for attachment_dict in attachments:
						attachment_name = attachment_dict["file_name"]
						file_extension = attachment_name.split(".")[-1]

						with col2:
							with st.container(border=True, key=str(uuid.uuid4())):
								if file_extension in image_extensions_list:
									attachment_image = attachment_dict.get("file")
									if attachment_image:
										attachment_image_decode_bytes = base64.b64decode(attachment_image)
										image = PIL.Image.open(BytesIO(attachment_image_decode_bytes))
										st.image(
											image,
											caption=attachment_name,
											use_container_width=True,
										)
								elif file_extension in docs_extensions_list:
									subcol1, subcol2 = st.columns([1, 3])
									with subcol1:
										st.image("assets/images/doc_icon.png", width=43)
									with subcol2:
										st.markdown(f"**{attachment_name}**")
								elif file_extension in video_extensions_list:
									subcol1, subcol2 = st.columns([1, 3])
									with subcol1:
										st.image(
											"assets/images/video_icon.jpg", width=43
										)
									with subcol2:
										st.markdown(f"**{attachment_name}**")
								elif file_extension in audio_extensions_list:
									subcol1, subcol2 = st.columns([1, 3])
									with subcol1:
										st.image(
											"assets/images/audio_icon.jpg", width=43
										)
									with subcol2:
										st.markdown(f"**{attachment_name}**")

				message(
					msg_dict["parts"][0]["text"], is_user=True, key=str(uuid.uuid4())
				)
				" "
			else:
				st.markdown(msg_dict["parts"][0]["text"])

	# "How can I help you today?" disappears on sending a message
	if st.session_state.new_chat:
		col1, col2, col3 = st.columns(3)
		with col2:
			st.subheader(translate_text("How can I help you today?", 80))

	# User Chat Input
	def on_submit_chat_input():
		st.session_state.user_can_send = False

	user_input = st.chat_input(
		translate_text("Ask Anything", 81),
		accept_file="multiple",
		file_type=[
			"png",
			"jpeg",
			"webp",
			"pdf",
			"txt",
			"aac",
			"flac",
			"mp3",
			"m4a",
			"mpeg",
			"mpga",
			"mp4",
			"opus",
			"pcm",
			"wav",
			"webm",
			"flv",
			"mov",
			"mpg",
			"wmv",
			"3gpp",
		],
		on_submit=on_submit_chat_input,
		disabled=not st.session_state.user_can_send,
	)

	if user_input:
		try:
			if st.session_state.new_chat:
				# Generating a new chat number
				chats_numbers = (
					users_ref.child(st.session_state.username)
					.child("history")
					.child("ai")
					.get()
				)
				chat_number = "1" if chats_numbers is None else str(len(chats_numbers) + 1)
				st.session_state.current_chat_num = chat_number
				st.session_state.new_chat = False

			contents = []  # Text and/or Files
			file_names = []
			image_files = []

			text_prompt = user_input.text
			uploaded_files = user_input.files

			# Displaying uploaded files and processing each uploaded file
			col1, col2 = st.columns([3, 1])

			for uploaded_file in uploaded_files:
				# Get the file's byte content
				file_bytes = uploaded_file.getvalue()
				# Base64 Encoding
				file_b64 = base64.b64encode(file_bytes).decode("utf-8")
				# Get the file's name
				file_name = uploaded_file.name
				# Get the file's MIME type
				file_mime_type = uploaded_file.type

				# Appending the file's name to the list of file names
				file_names.append(file_name)

				# Append the file as an inlineData part for the Gemini API
				contents.append({"mime_type": file_mime_type, "data": file_bytes})

				file_extension = file_name.split(".")[-1]
				with col2:
					with st.container(border=True, key=str(uuid.uuid4())):
						if file_extension in docs_extensions_list:
							subcol1, subcol2 = st.columns([1, 3])
							with subcol1:
								st.image("assets/images/doc_icon.png", width=43)
							with subcol2:
								st.markdown(f"**{file_name}**")
						elif file_extension in image_extensions_list:
							image_files.append(file_b64)
							st.image(
								uploaded_file, caption=file_name, use_container_width=True
							)
						elif file_extension in video_extensions_list:
							subcol1, subcol2 = st.columns([1, 3])
							with subcol1:
								st.image("assets/images/video_icon.jpg", width=43)
							with subcol2:
								st.markdown(f"**{file_name}**")
						elif file_extension in audio_extensions_list:
							subcol1, subcol2 = st.columns([1, 3])
							with subcol1:
								st.image("assets/images/audio_icon.jpg", width=43)
							with subcol2:
								st.markdown(f"**{file_name}**")

			message(text_prompt, is_user=True, key=str(uuid.uuid4()))
			" "

			prompt = prompt_chat_history(
				st.session_state.username, st.session_state.current_chat_num, text_prompt
			)
			contents.append(prompt)

			# Adding the user's message to the history
			add_chat_to_history(
				st.session_state.username,
				chat_num=st.session_state.current_chat_num,
				role="user",
				message=text_prompt,
				attached_file_names=file_names,
				attached_image_files=image_files,
			)

			placeholder = st.empty()  # dynamic container
			answer = ""

			for chunk in st.session_state.gemini_2_0_flash.generate_content(
				contents, stream=True
			):
				answer += chunk.text  # append new chunk
				placeholder.markdown(answer)

			# Enabling the user to send another message
			st.session_state.user_can_send = True

			# Adding the AI's message to the history
			add_chat_to_history(
				st.session_state.username,
				chat_num=st.session_state.current_chat_num,
				role="model",
				message=answer,
			)

			# Saving/Updating the chat history
			st.session_state.chat_history = (
				users_ref.child(st.session_state.username)
				.child("history")
				.child("ai")
				.child(f"chat{st.session_state.current_chat_num}")
				.child("data")
				.get()
				or []
			)

			chat_title_try = (
				users_ref.child(st.session_state.username)
				.child("history")
				.child("ai")
				.child(f"chat{st.session_state.current_chat_num}")
				.child("title")
				.get()
			)

			if len(st.session_state.chat_history) <= 2 and not isinstance(
				chat_title_try, str
			):
				generate_chat_title(
					st.session_state.username, st.session_state.current_chat_num
				)

			st.rerun()
		except Exception as e:
			st.session_state.user_can_send = True
			st.error("Something went wrong.")

# -----------------------------
# 📚 Learning Plans
# -----------------------------

elif "learning plans" in st.session_state.current_page.lower():
	st.title("📚 Create Your Learning Plan")
	
	# Check if user already has an active learning plan
	existing_plan = get_learning_plan(st.session_state.username, st.session_state.current_plan_num)
	
	if existing_plan and existing_plan.get("plan_data") and not st.session_state.make_new_plan:
		st.info("🎯 You already have an active learning plan! Visit the Learning Dashboard to continue your progress.")
		
		col1, col2 = st.columns(2)
		with col1:
			if st.button("Go to Dashboard", icon="📊", type="primary", use_container_width=True):
				st.session_state.current_page = "🏆 Learning Dashboard"
				st.rerun()
		
		with col2:
			if st.button("Create New Plan", icon="🆕", type="secondary", use_container_width=True):
				# Clear existing plan
				st.session_state.make_new_plan = True
				st.rerun()
		
		st.divider()
		
		# Show current plan details
		plan_data = existing_plan["plan_data"]
		st.subheader(f"Current Plan: {plan_data['plan_title']}")

		" "

		col1, col2 = st.columns(2)
		with col1:
			st.metric("Difficulty", plan_data['difficulty'])
		with col2:
			st.metric("Duration", f"{plan_data['total_days']} days")
		
		if plan_data['ar']['can_be_represented_in_3d']:
			st.write(f"🥽 **AR: {google_translate(plan_data['ar']['ar_search_term'].title(), st.session_state.language)}**")
		else:
			st.write("🥽 **AR: Not Applicable**")

		# Show first few days preview
		st.markdown("### 👀 Preview: First 3 Days")
		
		for day in plan_data["days"][:3]:
			with st.expander(f"Day {day['day']}: {day['title']}"):
				st.write(f"⏰ Estimated time: {day['estimated_time']} minutes")
				st.write("📖 *Content Preview:*")
				st.write(day['content'][:200] + "...")
				
				st.write(f"🎥 Video: {day['youtube_query']}")
				
				st.write(f"📝 {len(day['flashcards'])} flashcards")
				st.write(f"❓ {len(day['quiz'])} quiz questions")
		
		with stylable_container(
			key="ready_to_start_learning",
			css_styles="""
				{
					background-color: #4682B4;
					padding: 19px;
					border-radius: 10px;
				}
			"""
		):
			if st.button("Start Learning Now!", icon="🎯", type="tertiary"):
				st.session_state.current_page = "🏆 Learning Dashboard"
				st.session_state.start_learning = True
				st.rerun()
			
			" "

	else:
		# Create new learning plan form
		st.markdown("### 🎯 Let's create your personalized learning journey!")
		
		with st.form("learning_plan_form", border=True):
			# Topic input
			topic = st.text_input(
				"📖 What do you want to learn?",
				placeholder="e.g., Python Programming, World History, Biology, Spanish..."
			)
			
			# Difficulty selection
			difficulty = st.selectbox(
				"Choose your difficulty level",
				["Beginner", "Intermediate", "Advanced"],
				help="This will determine the complexity of content and exercises"
			)
			
			# Duration
			days = st.slider(
				"📅 How many days do you have?",
				min_value=3,
				max_value=30,
				value=7,
				help="We'll create a day-by-day plan that fits your timeline"
			)
			
			# Additional preferences
			st.markdown("#### 🔧 Additional Preferences (Optional)")
			
			col1, col2 = st.columns(2)
			
			with col1:
				include_ar = st.checkbox(
					"🥽 Include AR experiences when relevant",
					value=True,
					help="We'll add a 3D model for the topic that benefits in visual learning"
				)

			with col2:
				focus_area = st.selectbox(
					"🎯 Learning focus",
					["Balanced", "Theory Heavy", "Practice Heavy", "Visual Learning"],
					help="Customize your learning approach"
				)

				plan_language = st.selectbox(
					"🌐 Plan Language",
					[lang.title() for lang in list(languages_dict.keys())],
					index=27, # English
					help="All the material will be in this language"
				)
			
			custom_prompt = st.text_area(
				"📝 Custom Prompt (Optional)",
				height=150,
				placeholder="Specify a specific lenth of materials, Provide source websites, Cover a specific section, etc.",
			)
			
			# Submit button
			submitted = st.form_submit_button(
				"🚀 Generate My Learning Plan",
				type="primary",
				use_container_width=True
			)
			
			if submitted:
				if topic:
					with st.spinner("🤖 Creating your personalized learning plan..."):
						
						# Generate the learning plan
						plan_data = generate_learning_plan(topic, difficulty, days, st.session_state.age, plan_language, custom_prompt)
						
						if plan_data:
							# Update new plan status
							st.session_state.make_new_plan = False

							# Save to Firebase
							save_new_learning_plan(st.session_state.username, plan_data, include_ar)
							plans_dict = users_ref.child(st.session_state.username).child("history").child("learning_plans").get()
							st.session_state.current_plan_num = len(plans_dict)
							
							st.divider()
							st.success("🎉 Your learning plan has been created!")
							st.balloons()
							
							# Show plan overview
							st.markdown("### 📋 Plan Overview")
							
							col1, col2, col3 = st.columns(3)
							with col1:
								st.metric("📚 Title", plan_data["plan_title"])
							with col2:
								st.metric("⏱ Duration", f"{plan_data['total_days']} days")
							with col3:
								st.metric("🎚 Difficulty", plan_data["difficulty"])
							
							if plan_data['ar']['can_be_represented_in_3d']:
								st.write(f"🥽 **AR: {plan_data['ar']['ar_search_term'].title()}**")
							else:
								st.write("🥽 **AR: Not Applicable**")

							# Show first few days preview
							st.markdown("### 👀 Preview: First 3 Days")
							
							for day in plan_data["days"][:3]:
								with st.expander(f"Day {day['day']}: {day['title']}"):
									st.write(f"⏰ Estimated time: {day['estimated_time']} minutes")
									st.write("📖 *Content Preview:*")
									st.write(day['content'][:200] + "...")
									
									st.write(f"🎥 Video: {day['youtube_query']}")
									
									st.write(f"📝 {len(day['flashcards'])} flashcards")
									st.write(f"❓ {len(day['quiz'])} quiz questions")
							
							st.info("🏆 Go to your Learning Dashboard to continue your learning journey!")
							
						else:
							st.error("❌ Failed to generate learning plan. Please try again.")
				else:
					st.error("Please enter a topic you want to learn!")
		
		" "

		col1, col2 = st.columns(2)
		with col1:
			if st.button("Back", icon="🔙", type="secondary", use_container_width=True):
				st.session_state.make_new_plan = False
				st.rerun()
		with col2:
			if st.button("Go to Learning Dashboard", icon="📊", type="secondary", use_container_width=True):
				print("clicked")
				st.session_state.current_page = "🏆 Learning Dashboard"
				st.rerun()


# -----------------------------
# 🏆 Learning Dashboard
# -----------------------------

elif "learning dashboard" in st.session_state.current_page.lower():
	# Get user's learning plan
	full_plan = get_learning_plan(st.session_state.username, st.session_state.current_plan_num)
	
	if not full_plan or not full_plan.get("plan_data"):
		st.title("🏆 Learning Dashboard")
		st.info("📚 You don't have an active learning plan yet!")
		
		col1, col2 = st.columns(2)
		with col1:
			if st.button("📚 Create Learning Plan", type="primary", use_container_width=True):
				st.session_state.current_page = "📚 Learning Plans"
				st.session_state.make_new_plan = True
				st.rerun()
		
		with col2:
			if st.button("🏠 Back To Home", type="secondary", use_container_width=True):
				st.session_state.current_page = "🧠 Ask AI" # Back to Home
				st.rerun()
		
		st.stop()
	
	# Extract plan information
	learning_plan = full_plan["plan_data"]
	current_day = full_plan.get("current_day")
	completed_days = full_plan.get("completed_days", [])
	daily_progress = full_plan.get("daily_progress", {})
	plan_done = full_plan.get("completed", False)
	
	# Calculate progress
	total_days = learning_plan["total_days"]
	progress_percentage = ((int(current_day) - 1) / total_days) * 100
	streak = calculate_streak(st.session_state.username, st.session_state.current_plan_num)

	# Page title
	st.title("🏆 Learning Dashboard")

	if not st.session_state.start_learning or (st.session_state.start_learning and plan_done):
		# 1. Header with welcome, quote, and streak
		with stylable_container(
			key="header_card",
			css_styles="""
			{
				background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
				color: white;
				padding: 1.5rem;
				border-radius: 15px;
				margin-bottom: 1rem;
			}
			"""
		):
			col1, col2 = st.columns([2, 1])
			
			with col1:
				try:
					name = st.session_state.user_real_name.split(" ")[0]
				except:
					name = st.session_state.user_real_name

				st.markdown(f"### 👋 Welcome back, {name}!")
				st.markdown(f"{get_motivational_quote()}")
			
			with col2:
				st.markdown(f"<h3 style='text-align: center;'>🔥 Learning Streak</h3>", unsafe_allow_html=True)
				st.markdown(f"<h3 style='text-align: center;'>{streak} days</h3>", unsafe_allow_html=True)
			
			" "
		
		# 2. Progress Card
		with stylable_container(
			key="progress_card",
			css_styles="""
			{
				background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);
				color: #333333;
				padding: 1.5rem;
				border-radius: 15px;
				margin-bottom: 1rem;
			}
			"""
		):
			st.markdown("### 📈 Learning Progress")
			" "
			
			col1, _, col3 = st.columns([2, 0.5, 0.6])
			
			with col1:
				st.markdown(f"📚 Current Plan: **{learning_plan['plan_title']}**")
				st.progress(progress_percentage / 100)
				st.markdown(f"*Progress:* {current_day - 1}/{total_days} days completed ({progress_percentage:.0f}%)")
			
			with col3:
				if not plan_done:
					st.metric("Current Day", f"{current_day}/{total_days}")
					st.metric("Days Left", total_days - (current_day - 1))
		
		# 3. Smart Suggestion Card
		with stylable_container(
			key="suggestion_card",
			css_styles="""
			{
				background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%);
				color: #1a237e;
				padding: 1.5rem;
				border-radius: 15px;
				margin-bottom: 1rem;
			}
			"""
		):
			st.markdown("### 🎯 Smart Suggestions")
			" "
			
			if current_day <= total_days:
				next_day = learning_plan["days"][current_day - 1]
				estimated_time = next_day["estimated_time"]
				
				st.markdown(f"***Next Step:*** {next_day['title']}")
				st.markdown(f"***Estimated Time:*** {estimated_time} minutes")
				st.markdown("**Ready to continue your learning journey!**")
			else:
				st.markdown("🎉 *Congratulations!* You've completed your learning plan!")
				st.markdown("*Next Step:* Create a new learning plan or review your achievements")
		
		# 4. Mood Check-in Card
		with stylable_container(
			key="mood_card",
			css_styles="""
			{
				background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
				color: white;
				padding: 1.5rem;
				border-radius: 15px;
				margin-bottom: 1rem;
			}
			"""
		):
			st.markdown("### 😊 How are you feeling today?")
			" "
			
			col1, col2, col3 = st.columns(3)
			
			with col1:
				if st.button("😴 Tired", use_container_width=True):
					st.session_state.user_mood = "tired"
					st.info("💡 We'll suggest shorter study sessions with more breaks!")
			
			with col2:
				if st.button("🎯 Focused", use_container_width=True):
					st.session_state.user_mood = "focused"
					st.success("🚀 Great! You're ready for intensive learning!")
			
			with col3:
				if st.button("😰 Overwhelmed", use_container_width=True):
					st.session_state.user_mood = "overwhelmed"
					st.info("💝 Let's break things down into smaller, manageable chunks!")
		
		# 5. Badges Card
		badges = get_plan_badges(st.session_state.username, progress_percentage, st.session_state.current_plan_num)
		
		with stylable_container(
			key="badges_card",
			css_styles="""
			{
				background: linear-gradient(135deg, #ffd89b 0%, #19547b 100%);
				color: #2c2c2c;
				padding: 1.5rem;
				border-radius: 15px;
				margin-bottom: 1rem;
			}
			"""
		):
			st.markdown("### 🏅 Your Badges")
			" "
			
			if badges:
				cols = st.columns(min(len(badges), 4))
				for i, badge in enumerate(badges):
					with cols[i % 4]:
						st.markdown(f"""
						<div style="text-align: center; padding: 0.5rem;">
							<div style="font-size: 2rem;">{badge['icon']}</div>
							<div style="font-weight: bold; margin-top: 0.5rem;">{badge['name']}</div>
						</div>
						""", unsafe_allow_html=True)
			else:
				st.markdown("🎯 Complete your first day to start earning badges!")
		
		# 6. Quick Actions Card
		with stylable_container(
			key="actions_card",
			css_styles="""
			{
				background: linear-gradient(135deg, #ff9a9e 0%, #fecfef 100%);
				color: white;
				padding: 1.5rem;
				border-radius: 15px;
				margin-bottom: 1rem;
			}
			"""
		):
			st.markdown("### ⚡ Quick Actions")
			" "
			
			col1, col2 = st.columns(2)
			
			with col1:
				if not plan_done:
					if st.button("Continue Learning", icon="▶", type="primary", use_container_width=True):
						st.session_state.start_learning = True
						st.rerun()
				else:
					st.balloons()
					if st.button("Retake The Plan", icon="🔄", use_container_width=True):
						# Setting completed to False
						users_ref.child(st.session_state.username).child("history").child("learning_plans").child(f"plan{st.session_state.current_plan_num}").child("completed").set(False)

						# Resetting completed days and daily progress
						users_ref.child(st.session_state.username).child("history").child("learning_plans").child(f"plan{st.session_state.current_plan_num}").child("completed_days").set([])
						users_ref.child(st.session_state.username).child("history").child("learning_plans").child(f"plan{st.session_state.current_plan_num}").child("daily_progress").set({})

						# Setting current day to 1
						users_ref.child(st.session_state.username).child("history").child("learning_plans").child(f"plan{st.session_state.current_plan_num}").child("current_day").set(1)

						st.session_state.start_learning = True
						st.rerun()
			
			with col2:
				if st.button("New Learning Plan", icon="🆕", type="secondary", use_container_width=True):
					st.session_state.make_new_plan = True
					st.session_state.current_page = "📚 Learning Plans"
					st.rerun()

	# Current Lesson Display
	else:
		plans_dict = users_ref.child(st.session_state.username).child("history").child("learning_plans").get()
		if plans_dict:
			if st.session_state.current_plan_num:
				plan_num = st.session_state.current_plan_num
			else:
				plan_num = len(plans_dict)

		lesson_day = current_day
		current_lesson = learning_plan["days"][lesson_day - 1]
		
		st.divider()
		st.markdown(f"## 📚 Day {lesson_day}: {current_lesson['title']}")
		
		# Create tabs for different lesson components
		tab1, tab2, tab3, tab4, tab5 = st.tabs(["📖 Content", "🎥 Video", "🥽 AR Model", "🔄 Flashcards", "❓ Quiz"])

		# Content
		with tab1:
			st.markdown("### 📖 Reading Content")
			st.write(current_lesson['content'])
			" "

			if st.checkbox("Mark Content as Read", key=f"content_done_{lesson_day}", value=st.session_state.content_marked_done):
				st.session_state.content_marked_done = True
				update_daily_progress(st.session_state.username, lesson_day, "content", True)
		
		# Video
		with tab2:
			st.markdown("### 🎥 Recommended Video")
			
			# Youtube Video Embed
			youtube_queries = current_lesson['youtube_query'].split(",")
			video_ids = get_video_id(youtube_queries)
			
			for video_id in video_ids:
				components.html(f"""
					<div style="position: relative; padding-bottom: 56.25%; height: 0; overflow: hidden;">
						<iframe src="https://www.youtube.com/embed/{video_id}"
							style="position: absolute; top:0; left:0; width:100%; height:100%;"
							frameborder="0"
							allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture"
							allowfullscreen>
						</iframe>
					</div>
				""", height=600)
				" "

			if st.checkbox("Mark Video as Watched", key=f"video_done_{lesson_day}", value=st.session_state.video_marked_done):
				st.session_state.video_marked_done = True
				update_daily_progress(st.session_state.username, lesson_day, "video", True)
		
		# AR
		with tab3:
			if learning_plan['ar']['can_be_represented_in_3d']:
				st.markdown("### 🥽 AR Experience")
				
				if learning_plan['ar'].get('ar_data'):
					data = learning_plan['ar']['ar_data']

					components.html(
						f"""
						<div style="position: relative; padding-bottom: 56.25%; height: 0; overflow: hidden;">
							<iframe 
								title="Sketchfab Model"
								src="{data[1]}"
								frameborder="0"
								allowfullscreen 
								mozallowfullscreen="true"
								webkitallowfullscreen="true"
								style="position: absolute; top:0; left:0; width:100%; height:100%;">
							</iframe>
						</div>
						""",
						height=600
					)

					" "

					st.markdown(f"[👀 See In AR]({data[0]})")

					update_daily_progress(st.session_state.username, lesson_day, "ar", True)

				else:	
					if st.button("🚀 Generate AR Experience", key=f"ar_launch_{lesson_day}"):
						# Generate AR Experience
						with st.spinner("Generating Your AR Experience...", show_time=True):
							success, data = generate_ar(learning_plan['ar']['ar_search_term'])
							if success:
								save_ar_data(st.session_state.username, data, st.session_state.current_plan_num)
								st.rerun()
							else:
								st.error(f"Error generating your AR experience: {data}")

			else:
				st.info("🔍 AR experience not applicable for this topic")
		
		# Flashcards
		if f"flashcard_index_{lesson_day}" not in st.session_state:
			st.session_state[f"flashcard_index_{lesson_day}"] = 0

		current_index = st.session_state[f"flashcard_index_{lesson_day}"]

		with tab4:
			st.markdown("### 🔄 Flashcards")
			flashcards = current_lesson['flashcards']
			
			if flashcards:
				card = flashcards[current_index]
				
				with stylable_container(
					key=f"flashcard_{lesson_day}_{current_index}",
					css_styles="""
					{
						background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
						color: white;
						padding: 2rem;
						border-radius: 15px;
						text-align: center;
						margin: 1rem 0;
					}
					"""
				):
					st.markdown(f"<span style='color: black;'>**Card {current_index + 1} of {len(flashcards)}**</span>", unsafe_allow_html=True)
					st.write(" ")
					
					if f"show_answer_{lesson_day}_{current_index}" not in st.session_state:
						st.session_state[f"show_answer_{lesson_day}_{current_index}"] = False
					
					if not st.session_state[f"show_answer_{lesson_day}_{current_index}"]:
						st.subheader(f"*Question:* {card['question']}")
						if st.button("🔍 Show Answer", key=f"show_ans_{lesson_day}_{current_index}"):
							st.session_state[f"show_answer_{lesson_day}_{current_index}"] = True
							st.rerun()
					else:
						st.subheader(f"*Question:* {card['question']}")
						st.markdown(f"#### *Answer:* {card['answer']}")
						if st.button("❓ Hide Answer", key=f"hide_ans_{lesson_day}_{current_index}"):
							st.session_state[f"show_answer_{lesson_day}_{current_index}"] = False
							st.rerun()
				
				col1, _, col3 = st.columns([1, 1, 0.3])
				
				with col1:
					if current_index > 0:
						if st.button("⬅ Previous", key=f"prev_card_{lesson_day}"):
							st.session_state[f"flashcard_index_{lesson_day}"] -= 1
							st.rerun()

				with col3:
					if current_index < len(flashcards) - 1:
						if st.button("Next ➡", key=f"next_card_{lesson_day}"):
							st.session_state[f"flashcard_index_{lesson_day}"] += 1
							st.rerun()
				
				if current_index == len(flashcards) - 1 or st.session_state.flashcard_marked_done:
					st.session_state.flashcard_marked_done = True
					update_daily_progress(st.session_state.username, lesson_day, "flashcards", True)
					st.success("🎉 You've completed all flashcards for today")
		
		# Quiz
		with tab5:
			if not st.session_state.quiz_taken:
				st.markdown("### ❓ Daily Quiz")
			
			if f"quiz_started_{lesson_day}" not in st.session_state:
				st.session_state[f"quiz_started_{lesson_day}"] = False
			
			if st.session_state.quiz_taken and st.session_state.quiz_taken_data != []:
				quiz_results = st.session_state.quiz_taken_data
				result1, result2 = quiz_results
				score = (result1 / result2) * 100
						
				st.markdown("### 📊 Quiz Results")
				st.metric("Score", f"{score:.0f}%")
				st.progress(score / 100)
				
				if score >= 70:
					st.success(f"🎉 Great job! You scored {result1}/{result2}")
					
					# Mark day as complete
					completed_days_list = learning_plan.get("completed_days", [])
					today_str = str(datetime.date.today())
					completed_days_list.append(today_str)

					update_daily_progress(st.session_state.username, lesson_day, "quiz", True)
					
					next_day_btn_text = "🎯 Continue to Next Day"
					if lesson_day >= total_days:
						next_day_btn_text = "🎉 Finish Learning Plan"

					if st.button(next_day_btn_text):
						if not lesson_day > total_days:
							update_current_day(st.session_state.username, completed_days_list, lesson_day + 1, st.session_state.current_plan_num)

						# Setting marked as done session states to False
						st.session_state.content_marked_done = False
						st.session_state.video_marked_done = False
						st.session_state.flashcard_marked_done = False
						st.session_state.quiz_taken = False
						st.session_state.quiz_taken_data = []
						st.session_state.start_quiz = False
						current_index = 0

						if lesson_day >= total_days:
							users_ref.child(st.session_state.username).child("history").child("learning_plans").child(f"plan{st.session_state.current_plan_num}").child("completed").set(True)
							st.session_state.start_learning = False
							st.balloons()
							st.success("🎉 Congratulations! You've completed your entire learning plan!")

						st.rerun()
				
				else:
					st.warning(f"📚 You scored {result1}/{result2}. Review the material and try again!")
					st.info("💡 Tip: Go back and review the content, video, and flashcards before retaking the quiz.")

				" "
				if st.button("🚀 Retake Quiz"):
					st.session_state.quiz_taken = False
					st.session_state.quiz_taken_data = []
					
					st.session_state.start_quiz = True
					st.session_state[f"quiz_started_{lesson_day}"] = True
					st.session_state[f"quiz_answers_{lesson_day}"] = {}
					st.rerun()

			elif not st.session_state[f"quiz_started_{lesson_day}"] or not st.session_state.start_quiz:
				st.write("Ready to test your knowledge?")
				if st.button("🚀 Start Quiz", key=f"start_quiz_{lesson_day}", type="primary"):
					st.session_state.start_quiz = True
					st.session_state[f"quiz_started_{lesson_day}"] = True
					st.session_state[f"quiz_answers_{lesson_day}"] = {}
					st.rerun()
			
			else:
				quiz_questions = current_lesson['quiz']
				
				with st.container(border=True):
					user_answers = {}
					
					for i, question in enumerate(quiz_questions):
						st.markdown(f"*Question {i+1}:* {question['question']}")
						user_answers[i] = st.radio(
							"Select your answer:",
							question['options'],
							key=f"quiz_q_{lesson_day}_{i}"
						)
					
					" "
					if st.button("📝 Submit Quiz", type="primary", use_container_width=True):
						st.session_state.quiz_taken = True

						# Calculate score
						correct_answers = 0
						for i, question in enumerate(quiz_questions):
							if question['options'].index(user_answers[i]) == question['correct']:
								correct_answers += 1
						
						st.session_state.quiz_taken_data.append(correct_answers)
						st.session_state.quiz_taken_data.append(len(quiz_questions))
						st.rerun()

		# Back to dashboard button
		st.divider()

		col1, col2 = st.columns(2)
		with col1:
			if st.button("Back to Dashboard", icon="📊", use_container_width=True):
				st.session_state.start_learning = False
				st.rerun()
		with col2:
			if st.button("Back to Learning Plans", icon="📚", use_container_width=True):
				st.session_state.start_learning = False
				st.session_state.current_page = "📚 Learning Plans"
				st.rerun()

# -----------------------------
# AR, Content and Quiz
# -----------------------------
elif "ar, content and quiz package" in st.session_state.current_page.lower():

	# Pages Titles with checkboxes
	titles_text = []
	ar_checkbox = st.checkbox("Learn with AR", value=True)
	content_checkbox = st.checkbox("Content Generation", value=True)
	quiz_checkbox = st.checkbox("Quiz Generation", value=True)

	st.divider()
	
	if ar_checkbox:
		titles_text.append("Learn with AR")
	if content_checkbox:
		titles_text.append("Content Generation")
	if quiz_checkbox:
		titles_text.append("Quiz Generation")

	def replace_last_comma(text: str) -> str:
		# Split by commas
		parts = text.rsplit(",", 1)  # split only at the last comma
		if len(parts) == 2:
			return parts[0] + " and" + parts[1]
		return text  # if there's no comma, just return as is

	st.header(replace_last_comma(", ".join(titles_text)))
	" "

	if ar_checkbox or content_checkbox or quiz_checkbox:
		# Required fields
		topic_name = st.text_input("Topic Name", placeholder="What are you learning about today?")

		# Content Section
		if content_checkbox:
			st.divider()
			grade = st.text_input("Grade (Optional)", placeholder="E.g. Grade 7, Third Prep")
			content_language = st.text_input("Content Language", value="English")
			content_length = st.select_slider("Content Length", options=["Auto", "Short", "Medium", "Long", "Extra Long"], value="Auto")
			content_complexity = st.select_slider("Content Complexity", options=["Auto", "Easy", "Medium", "Hard", "Extra Hard"], value="Auto")
			custom_prompt = st.text_area("Custom Prompt (Optional)", placeholder="Do you need to learn about a specific thing in particular? Or want the content from a specific source?", height=250)

		# Generate Button
		st.divider()
		if st.button("Generate"):
			if topic_name:
				" "
				model_embed_url = None

				if ar_checkbox:
					with st.spinner("Generating Your AR Experience...", show_time=True):
						success, data = generate_ar(topic_name)

					with st.expander("AR Results"):
						if success:
							# The AR Embed
							model_embed_url = data[1]

							components.html(
								f"""
								<div style="position: relative; padding-bottom: 56.25%; height: 0; overflow: hidden;">
									<iframe 
										title="Sketchfab Model"
										src="{model_embed_url}"
										frameborder="0"
										allowfullscreen 
										mozallowfullscreen="true"
										webkitallowfullscreen="true"
										style="position: absolute; top:0; left:0; width:100%; height:100%;">
									</iframe>
								</div>
								""",
								height=600
							)
							st.write(f"**Live Your [AR Experience]({data[0]})**")
							st.divider()

							# QR Code Creation
							st.write("Or Scan The QR Code 👇") 
							# Creating the QR code
							qr_code = qrcode.make(data[0])
							# Saving the QR code to a buffer
							qr_code_buffer = BytesIO()
							qr_code.save(qr_code_buffer, "PNG")
							# Rewend the buffer to the beginning
							qr_code_buffer.seek(0)

							st.image(qr_code_buffer, width=230)

						else:
							st.error(data)
				
				if content_checkbox:
					with st.spinner("Generating Your Article and Document...", show_time=True):
						content_results = generate_content(
							topic_name,
							age=st.session_state.age,
							grade=grade,
							language=content_language,
							length=content_length,
							complexity=content_complexity,
							custom_prompt=custom_prompt,
							model_embed_url=model_embed_url
						)
					
					with st.expander("Content Results"):
						if isinstance(content_results, list):
							st.write(f"**Go Inside Your [Article]({content_results[0]})**")

							st.divider()
							st.write("Or Scan The QR Code 👇")
							st.image(content_results[1], width=230)

							st.divider()
							st.download_button(
								"Download Content",
								data=content_results[2],
								file_name=f"{topic_name}.docx",
								mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
							)
						else:
							st.error(content_results)
			else:
				st.error("Please enter the topic you are learning about.")


elif "learn with ar" in st.session_state.current_page.lower():
	# Language selection
	langs_dict = GoogleTranslator().get_supported_languages(as_dict=True)
	langs_selection_list = [lang.title() for lang in list(langs_dict.keys())]

	st.session_state.target_language = st.selectbox(
		"Select Language",
		[lang.title() for lang in langs_selection_list],
		index=langs_selection_list.index("English"),
	)
	target_language_name = st.session_state.target_language
	st.session_state.target_language = langs_dict[
		st.session_state.target_language.lower()
	]

	st.divider()
	# --- Home Page Layout ---
	# Logo
	col1, col2, col3 = st.columns([1.8, 2, 1])

	with col2:
		st.image(
			"assets/images/logo.png", width=325, caption="Learn with Agumented Reality"
		)

	" "
	st.divider()

	def automated_processes(
		search_query: str,
		age: int,
	):
		st.session_state.backend_running = True
		if st.session_state.history_rerun:
			st.session_state.history_rerun = False
			st.rerun()

		if "ar_done" not in st.session_state:
			st.session_state.ar_done = False
		if "content_done" not in st.session_state:
			st.session_state.content_done = False
		if "quiz_done" not in st.session_state:
			st.session_state.quiz_done = False
		if "echo3d_time" not in st.session_state:
			st.session_state.echo3d_time = 0
		if "ar_error" not in st.session_state:
			st.session_state.ar_error = False
		if "content_error" not in st.session_state:
			st.session_state.content_error = False
		if "quiz_error" not in st.session_state:
			st.session_state.quiz_error = False

		if "ar_skip" in st.session_state and st.session_state.ar_skip is None:
			st.session_state.ar_skip = False
		elif "ar_skip" not in st.session_state:
			st.session_state.ar_skip = False

		SKETCHFAB_API_TOKEN = st.secrets["SKETCHFAB_TOKEN"]

		# Cleaning the model name to make it file-safe
		model_name = re.sub(r'[\\/*?:"<>|]', "_", search_query)

		# Model Name to English
		model_name = GoogleTranslator(source="auto", target="en").translate(model_name)

		# --- The part where the user chooses a 3d model from 15 ---
		first_show_models_num = 9
		increment = 6
		if not st.session_state.is_pc:
			first_show_models_num -= 3
			increment -= 3

		num_results = first_show_models_num + increment

		if "model_uid" not in st.session_state:
			st.session_state.model_uid = None

		if st.session_state.model_uid is None:

			# --- Process 1: Download from Sketchfab unless skipping AR ---

			# Step 1.1: Search for 3d models on Sketchfab
			def search_models(query, limit):
				url = "https://api.sketchfab.com/v3/search"
				authorization_header = {"Authorization": f"Token {SKETCHFAB_API_TOKEN}"}
				params = {
					"q": query,
					"type": "models",
					"downloadable": True,
					"count": limit,
				}
				response = requests.get(
					url, headers=authorization_header, params=params
				)
				response.raise_for_status()
				return response.json()["results"]

			# Step 1.2: Asking the user to select a model from 15, skip the AR or cancel the whole process
			st.title("🎯 " + translate_text("Select a 3D Model"))

			with st.spinner(translate_text("Searching Sketchfab...")):
				models = search_models(model_name, num_results)
				first_show = models[:first_show_models_num]

			if not models:
				st.error(translate_text("No models found. Skipping AR."))
				if st.button(translate_text("Continue")):
					st.session_state.model_uid = "no models found"
					st.rerun()
				if st.button(translate_text("Cancel")):
					st.session_state.model_uid = "cancel process"
					st.rerun()

			elif st.session_state.is_pc:
				# Desktop Layout

				# Displaying the first 9 models' thumbnails
				cols = st.columns(3)
				for i, model in enumerate(first_show):
					with cols[i % 3]:
						st.image(
							model["thumbnails"]["images"][0]["url"],
							caption=model["name"],
							width=425,
						)
						if st.button(label=f"Select #{i+1}", key=f"select_{i}"):
							st.session_state.model_uid = model["uid"]
							st.session_state.history_model_thumbs.append(
								model["thumbnails"]["images"][0]["url"]
							)
							st.rerun()

				# View more expander button
				" "
				with st.expander(translate_text("View more")):
					second_show = models[first_show_models_num:]
					expander_cols = st.columns(3)
					for i, model in enumerate(second_show):
						i += first_show_models_num
						with expander_cols[i % 3]:
							st.image(
								model["thumbnails"]["images"][0]["url"],
								caption=model["name"],
								width=425,
							)
							if st.button(label=f"Select #{i+1}", key=f"select_{i}"):
								st.session_state.model_uid = model["uid"]
								st.session_state.history_model_thumbs.append(
									model["thumbnails"]["images"][0]["url"]
								)
								st.rerun()

				col1, col2 = st.columns([6, 1])

				# Cancel button to stop the process
				with col1:
					if st.button("🛑 " + translate_text("Cancel")):
						st.session_state.model_uid = "cancel process"
						st.rerun()

				# Skip AR button to skip AR and continue with the content and quiz
				with col2:
					if st.button("⏭️ " + translate_text("Skip AR")):
						st.session_state.model_uid = "skip ar"
						st.rerun()
			else:
				# 📱 Mobile layout: show all models in one vertical column
				for i, model in enumerate(first_show):
					st.image(
						model["thumbnails"]["images"][0]["url"],
						caption=model["name"],
						width=300,  # smaller width for mobile
					)
					if st.button(label=f"Select #{i+1}", key=f"select_{i}"):
						st.session_state.model_uid = model["uid"]
						st.session_state.history_model_thumbs.append(
							model["thumbnails"]["images"][0]["url"]
						)
						st.rerun()

				# View more expander for mobile
				with st.expander(translate_text("View more")):
					second_show = models[first_show_models_num:]
					for i, model in enumerate(second_show):
						i += first_show_models_num
						st.image(
							model["thumbnails"]["images"][0]["url"],
							caption=model["name"],
							width=300,
						)
						if st.button(label=f"Select #{i+1}", key=f"select_{i}"):
							st.session_state.model_uid = model["uid"]
							st.session_state.history_model_thumbs.append(
								model["thumbnails"]["images"][0]["url"]
							)
							st.rerun()

				# Cancel and Skip AR buttons (mobile version: vertical)
				" "
				if st.button("🛑 " + translate_text("Cancel")):
					st.session_state.model_uid = "cancel process"
					st.rerun()

				if st.button("⏭️ " + translate_text("Skip AR")):
					st.session_state.model_uid = "skip ar"
					st.rerun()

				st.warning(
					"Please select a model or choose to cancel/skip to continue."
				)

		# After the user decides his choice 👇
		elif st.session_state.model_uid is not None:

			" "

			if st.session_state.model_uid == "cancel process":
				st.session_state.model_name = None
				st.session_state.age = None
				st.session_state.model_uid = None
				st.session_state.user_content_length = None
				st.session_state.user_content_complexity = None
				st.session_state.user_custom_prompt = None

				st.session_state.ar_skip = None
				st.session_state.ar_done = False
				st.session_state.content_done = False
				st.session_state.quiz_done = False

				st.session_state.echo3d_time = 0
				st.session_state.telegraph_time = 0
				st.session_state.quiz_time = 0

				st.session_state.ar_error = False
				st.session_state.content_error = False
				st.session_state.quiz_error = False

				st.session_state.backend_running = False
				st.session_state.history_rerun = True
				st.session_state.history_rerun2 = True

				st.rerun()
			elif st.session_state.model_uid == "skip ar":
				st.session_state.ar_skip = True
			else:
				model_uid = st.session_state.model_uid

			# Creating the tabs to display the 3 outputs
			ar_tab, content_tab, quiz_tab = st.tabs(
				["AR", translate_text("Content"), translate_text("Quiz")]
			)

			if st.session_state.model_uid == "no models found":
				with ar_tab:
					st.error(translate_text("No models found. Skipping AR."))

			elif st.session_state.ar_skip and st.session_state.ar_error:
				with ar_tab:
					st.error(
						"⚠️ "
						+ translate_text(
							"An error occured while generating your AR experience"
						)
					)
			elif st.session_state.ar_skip:
				with ar_tab:
					st.info("🔃 " + translate_text("AR is abandoned"))

			# If the user chose a model
			if not st.session_state.ar_done and not st.session_state.ar_skip:

				with st.spinner(translate_text("Creating your AR Experience...")):
					sketchfab_download_start_time = time.time()

					with st.spinner(translate_text("Downloading the 3D model...")):

						# Step 1.3: Setting up the download API request
						headers = {"Authorization": f"Token {SKETCHFAB_API_TOKEN}"}
						download_url = (
							f"https://api.sketchfab.com/v3/models/{model_uid}/download"
						)
						download_response = requests.get(download_url, headers=headers)

						if download_response.status_code != 200:
							st.session_state.ar_error = True
							st.session_state.ar_skip = True
							st.rerun()

						glb_url = download_response.json().get("glb", {}).get("url")
						if not glb_url:
							st.session_state.ar_error = True
							st.session_state.ar_skip = True
							st.rerun()

						# Step 1.4: Downloading the model in GLB and holding it in memory
						file_response = requests.get(glb_url)
						if file_response.status_code != 200:
							st.session_state.ar_error = True
							st.session_state.ar_skip = True
							st.rerun()

						glb_binary_data = file_response.content

					sketchfab_download_end_time = time.time()

					# --- Process 2: Creating an AR project on echo3D ---
					echo3d_start_time = time.time()

					# Step 2.1: Load Secret Variables
					echo3d_api_key = st.secrets["ECHO3D_API_KEY"]
					echo3d_email = st.secrets["ECHO3D_EMAIL"]
					echo3d_user_key = st.secrets["ECHO3D_User_Authentication_Key"]
					echo3d_security_key = st.secrets["ECHO3D_SECURITY_KEY"]

					# Step 2.3: Setting up the request payload giving the GLB file and the other parameters
					files = {"file_model": (f"{model_name}.glb", glb_binary_data)}
					data = {
						"key": echo3d_api_key,
						"email": echo3d_email,
						"userKey": echo3d_user_key,
						"target_type": 2,  # Have 0 in a quick video to show the judges
						"hologram_type": 2,
						"secKey": echo3d_security_key,
						"type": "upload",
					}

					# Step 2.4: Send the request to upload the model to Echo3D
					try:
						upload_resp = requests.post(
							"https://api.echo3D.com/upload", data=data, files=files
						)
						upload_resp.raise_for_status()
					except Exception as e:
						st.write(f"Error on step 2.4: {e}")
						st.session_state.ar_error = True
						st.session_state.ar_skip = True
						st.rerun()

					result = upload_resp.json()

					# Step 2.5: Saving the results
					ar_url = result.get("additionalData", {}).get("shortURL")

					echo3d_end_time = time.time()
					if ar_url:
						# Generating the QR Code and saving it to a buffer in bytes
						echo3d_ar_qr = qrcode.make(ar_url)
						echo3d_ar_qr_buffer = BytesIO()
						echo3d_ar_qr.save(
							echo3d_ar_qr_buffer, format="PNG"
						)  # Save as PNG to buffer
						echo3d_ar_qr_buffer.seek(
							0
						)  # Rewind the buffer (to read from the beginning)

						st.session_state.ar_url = ar_url
						st.session_state.echo3d_ar_qr_buffer = echo3d_ar_qr_buffer
						st.session_state.echo3d_time = (
							(echo3d_end_time - echo3d_start_time)
							- (
								sketchfab_download_end_time
								- sketchfab_download_start_time
							)
							- 3
						)

						st.session_state.history_ar_links.append(
							st.session_state.ar_url
						)
						st.session_state.history_ar_qrs.append(
							st.session_state.echo3d_ar_qr_buffer
						)

						st.session_state.ar_done = True

					else:
						st.session_state.ar_error = True
						st.session_state.ar_skip = True
						st.rerun()

					# Step 2.6: Displaying the results
					with ar_tab:
						st.write(
							f"➡️  **{translate_text('AR Link:')} {st.session_state.ar_url}**"
						)

						st.image(
							st.session_state.echo3d_ar_qr_buffer,
							caption=translate_text("Scan with your phone") + " 📱",
							width=250,
						)
						st.divider()
						st.write(
							translate_text(
								f"Made in {st.session_state.echo3d_time:.1f} seconds"
							)
						)

			elif st.session_state.ar_done and not st.session_state.ar_skip:
				with ar_tab:
					st.write(
						f"➡️  **{translate_text('AR Link:')} {st.session_state.ar_url}**"
					)

					st.image(
						st.session_state.echo3d_ar_qr_buffer,
						caption=translate_text("Scan with your phone") + " 📱",
						width=250,
					)
					st.divider()
					st.write(
						translate_text(
							f"Made in {st.session_state.echo3d_time:.1f} seconds"
						)
					)

			# HTML cleaner function
			def strip_html(html):
				soup = BeautifulSoup(html, "html.parser")
				return soup.get_text()

			if (
				"content_done" in st.session_state
				and not st.session_state.content_done
				and not st.session_state.content_error
			):
				# Adding the model name to the history
				st.session_state.history_model_names.append(model_name)

				# --- Process 3: Generating Content and Publishing on Telegraph ---
				with st.spinner(translate_text("Generating the content...")):
					telegraph_start_time = time.time()

					# Step 3.1: Getting some images to use from unsplash to show in the content
					UNSPLASH_ACCESS_KEY = st.secrets["UNSPLASH_ACCESS_KEY"]

					# Function to get 2 images from Unsplash
					def get_images_from_unsplash(model_name, count=2):
						url = "https://api.unsplash.com/search/photos"

						# Headers for the API request (with your key)
						headers = {"Authorization": f"Client-ID {UNSPLASH_ACCESS_KEY}"}

						# Search parameters
						params = {"query": model_name, "per_page": count}

						# Send the request
						response = requests.get(url, headers=headers, params=params)
						data = response.json()

						# Get the image URLs from the results
						images = []

						if len(data["results"]) == 0:
							return None

						for result in data.get("results", []):
							images.append(result["urls"]["regular"])

						return images

					# Calling the function and getting the result
					image_urls = get_images_from_unsplash(model_name)

					if st.session_state.ar_skip:
						st.session_state.history_model_thumbs.append(image_urls[0])

					# Convert the list of image URLs to a comma-separated string
					if image_urls:
						images = ", ".join(image_urls)

					# Step 3.2: Generating Content using Gemini API

					telegraph_prompt = f"""You are an AI API in an app which generates content about given input suitable for the given age in terms of words complexity, content length, etc.
				Strict Notes: You don't return content, but HTML code with the content; you must only return the HTML code (no CSS, JS, or any other text).
				and you must only write the HTML code with the Available tags: a, aside, b, blockquote, br, code, em, figcaption, figure, h3, h4, hr, i, iframe, img, li, ol, p, pre, s, strong, u, ul, video. (only them) to fit telegraph syntax. You must not ever use any tag other than these even <html> and <body> tags.
				Note: Please don't make empty bullet points. And make the title of the whole page in the required language.
				At the end, it is preferable to add a 'Did you know' section as well as a 'What is next?' section, including 'If you learned about ...., so you might also like: .....' with real links to other articles.
				Here are images to use: {images}
				Content Language: {target_language_name} (where you only change the text inside the HTML tags, not the HTML tags themselves)
				Generate content for a student (user) at age {st.session_state.age} 
				"""
					if st.session_state.user_grade:
						telegraph_prompt += f"(in grade: {st.session_state.user_grade})"

					if (
						st.session_state.user_content_length
						and st.session_state.user_content_length.lower() != "auto"
					):
						telegraph_prompt += f"The requested content length is: {st.session_state.user_content_length}. Stick to it."
					elif st.session_state.user_content_length:
						telegraph_prompt += (
							f"Determine the content length based on the age"
						)

					if (
						st.session_state.user_content_complexity
						and st.session_state.user_content_complexity.lower() != "auto"
					):
						telegraph_prompt += f"The requested content complexity is: {st.session_state.user_content_complexity}. Stick to it."
					elif st.session_state.user_content_complexity:
						telegraph_prompt += (
							f"Determine the content complexity based on the age"
						)

					if st.session_state.user_custom_prompt:
						telegraph_prompt += f"Here is a custom prompt from the user: {st.session_state.user_custom_prompt}. If it is not relevant or meaningful, ignore it. Ignore it like he didn't say it, don't even mention it."

					# Send the prompt to the model to generate content
					try:
						response = st.session_state.gemini_2_0_flash.generate_content(
							telegraph_prompt
						)
						html_code = re.sub(
							r"^```html\s*|```$",
							"",
							response.text.strip(),
							flags=re.MULTILINE,
						)
					except:
						st.session_state.content_error = True
						st.session_state.history_content_links.append("failed")
						st.session_state.history_content_qrs.append("failed")
						st.session_state.history_content_docs.append("failed")
						st.rerun()

					# Cleaning the HTML code
					def telegraph_html_cleaning(html_code):
						soup = BeautifulSoup(html_code, "html.parser")
						allowed_tags = [
							"a",
							"aside",
							"b",
							"blockquote",
							"br",
							"code",
							"em",
							"figcaption",
							"figure",
							"h3",
							"h4",
							"hr",
							"i",
							"iframe",
							"img",
							"li",
							"ol",
							"p",
							"pre",
							"s",
							"strong",
							"u",
							"ul",
							"video",
						]
						for tag in soup.find_all(True):  # True = all tags
							if tag.name not in allowed_tags:
								tag.decompose()  # removes the tag and all its content

						return str(soup)

					html_code = telegraph_html_cleaning(html_code)

					# Step 3.3: Publishing on Telegraph
					telegraph_access_token = st.secrets["TELEGRAPH_ACCESS_TOKEN"]
					telegraph = Telegraph(access_token=telegraph_access_token)

					response = telegraph.create_page(
						title=translate_text(
							f"{translate_text(model_name.title())} for Age {age}"
						),
						html_content=html_code,
					)

					content_url = "https://telegra.ph/" + response["path"]

					# Step 3.4: Generating a word document with the content
					st.session_state.clean_passage = strip_html(html_code)

					# Create the Word document
					content_doc = Document()

					# Add a centered bold title
					content_doc_heading = content_doc.add_heading(level=1)
					content_doc_run = content_doc_heading.add_run(
						translate_text(f"{model_name.title()} for Age {age}")
					)
					content_doc_run.font.name = "Arial"
					content_doc_run.font.size = Pt(20)
					content_doc_run.bold = True
					content_doc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

					# Add spacing
					content_doc.add_paragraph("")

					# Add the full cleaned passage, paragraph by paragraph
					for line in st.session_state.clean_passage.strip().split("\n"):
						if line.strip():  # Skip empty lines
							content_doc_para = content_doc.add_paragraph()
							content_doc_run = content_doc_para.add_run(line.strip())
							content_doc_run.font.name = "Arial"
							content_doc_run.font.size = Pt(16)
							if is_rtl():
								content_doc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
							else:
								content_doc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

					# Save to memory (BytesIO)
					content_doc_buffer = BytesIO()
					content_doc.save(content_doc_buffer)
					content_doc_buffer.seek(0)

					# Step 3.5: Generating QR Code
					telegraph_qr = qrcode.make(content_url)
					telegraph_qr_buffer = BytesIO()
					telegraph_qr.save(
						telegraph_qr_buffer, format="PNG"
					)  # Save as PNG to buffer
					telegraph_qr_buffer.seek(0)
					telegraph_end_time = time.time()

					# Step 3.6: Saving results
					st.session_state.content_url = content_url
					st.session_state.telegraph_qr_buffer = telegraph_qr_buffer
					st.session_state.content_doc_buffer = content_doc_buffer
					st.session_state.telegraph_time = (
						telegraph_end_time - telegraph_start_time
					)

					st.session_state.history_content_links.append(
						st.session_state.content_url
					)
					st.session_state.history_content_qrs.append(
						st.session_state.telegraph_qr_buffer
					)
					st.session_state.history_content_docs.append(
						st.session_state.content_doc_buffer
					)

					st.session_state.content_done = True

					# Step 3.7: Displaying results
					with content_tab:
						st.write(
							f"➡️  **{translate_text("Studying Content Link")}: {st.session_state.content_url}**"
						)

						st.image(
							st.session_state.telegraph_qr_buffer,
							caption=translate_text("Scan to open the content") + " 📱",
							width=250,
						)
						st.divider()
						st.download_button(
							label="📄 " + translate_text("Download Content Document"),
							data=st.session_state.content_doc_buffer,
							file_name=f"{model_name}_content.docx",
							mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
						)
						st.divider()
						st.write(
							translate_text(
								f"Made in {st.session_state.telegraph_time:.1f} seconds"
							)
						)

			elif st.session_state.content_done and not st.session_state.content_error:
				with content_tab:
					st.write(
						f"➡️  **{translate_text("Studying Content Link")}: {st.session_state.content_url}**"
					)

					st.image(
						st.session_state.telegraph_qr_buffer,
						caption=translate_text("Scan to open the content") + " 📱",
						width=250,
					)
					st.divider()
					st.download_button(
						label="📄 " + translate_text("Download Content Document"),
						data=st.session_state.content_doc_buffer,
						file_name=f"{model_name}_content.docx",
						mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
					)
					st.divider()
					st.write(
						translate_text(
							f"Made in {st.session_state.telegraph_time:.1f} seconds"
						)
					)

			if (
				"quiz_done" in st.session_state
				and not st.session_state.quiz_done
				and not st.session_state.quiz_error
			):
				# --- Process 4: Generating a Quiz on the Content Above ---
				with st.spinner(translate_text("Generating the Quiz...")):
					quiz_start_time = time.time()

					# Step 4.1: Generating a response from Gemini with declaring a strict form in the response to extract quiz data without errors
					questions_num = 10

					quiz_generation_prompt = f"""
					You will be given a passage and you should give me questions with their choices and correct answer of {questions_num} questions exactly in this form:
					Quiz title: ... (Don't write 'Quiz' or 'Form' at the end)
					Question type: multiple choice / multiple select / true or false
					Question name: Q1, 2, 3, ... : like Question title  (only alphanumerics, _) (eg: Q3_The_Capital_of_Egypt)
					Question text: ......
					Choices:
					Choice A: ...
					Choice B: ...
					Choice C: ...
					Choice D: ...
					Correct Choice: Choice choice_letter / Choice letter, Choice letter  (eg: Choice A / Choice B, Choice C)

					Notes:
					1. You must Only return my desired form (no more text or explaination), that is strict!
					2. You must write these starting words in English and as they are.
					2. You can increase the choices up to 10 choices if needed while it is usually 4.
					3. Make your own ratio to use the 3 types of question.

					Passage:
					{st.session_state.clean_passage}
					"""

					gemini_quiz_data = st.session_state.gemini_2_0_flash.generate_content(
						quiz_generation_prompt
					).text

					TYPEFORM_API_KEY = st.secrets["TYPEFORM_API_TOKEN"]
					headers = {
						"Authorization": f"Bearer {TYPEFORM_API_KEY}",
						"Content-Type": "application/json",
					}

					# Step 4.2: Capturing and Saving Quiz Data
					lines_list = gemini_quiz_data.strip().splitlines()

					questions_types = []
					questions_refs = []
					questions_texts = []
					questions_choices = []
					correct_choices = []
					st.session_state.quiz_title = "TeachAR Quiz"  # A default name
					current_question_list = []

					current_question_num = 0
					for line in lines_list:
						line = line.strip()

						if line.lower().startswith("quiz title"):
							current_quiz_title = line.split(":", 1)[1].strip()
							st.session_state.quiz_title_nt = current_quiz_title
							st.session_state.quiz_title = translate_text(
								current_quiz_title
							)

						if line.lower().startswith("question type"):
							questions_types.append(line.split(":", 1)[1].strip())

						if line.lower().startswith("question name"):
							questions_refs.append(
								line[len("question name: ") :].strip()
							)

						if line.lower().startswith("question text"):
							question_text = translate_text(
								line[len("question text: ") :].strip()
							)
							questions_texts.append(question_text)

						if line.lower().startswith("choice "):
							choice_text = translate_text(line.split(":", 1)[1].strip())
							current_question_list.append(choice_text)

						if line.lower().startswith("correct choice"):
							current_question_num += 1
							split = line.split(":", 1)
							correct_choices_letters = []

							if "," in split[1]:
								correct_choices_num = len(split[1].split(","))
								for i in range(correct_choices_num):
									letter = (
										split[1]
										.strip()
										.split(",")[i]
										.strip()
										.split(" ")[1]
										.strip()
									)
									correct_choices_letters.append(letter)
							else:
								correct_choices_letters.append(split[1].strip()[7])
							correct_choices.append(correct_choices_letters)

							questions_choices.append(
								{f"Q{current_question_num}": current_question_list}
							)
							current_question_list = []

					questions_refs = [
						ref.replace(":", "").replace(" ", "_") for ref in questions_refs
					]  # One more step to clean question references to match Typeform ID rules

					# Step 4.3: Setting up the Payload and sending a request
					# Creating the fields paramater by using the extracted quiz data
					fields = []
					for qtype, qref, qtext, choice_dict, q_correct in zip(
						questions_types,
						questions_refs,
						questions_texts,
						questions_choices,
						correct_choices,
					):
						tf_type = "multiple_choice"
						field = {
							"ref": qref,
							"type": tf_type,
							"title": qtext,
							"properties": {},
						}

						# Build choices list
						choices = []
						for choices_lst in choice_dict.values():
							for label, i in zip(choices_lst, range(len(choices_lst))):
								ref = f"{qref[:2]}_choice_{i + 1}".lower()
								appended_text = {"label": f"{label}", "ref": ref}
								choices.append(appended_text)

						field["properties"]["choices"] = choices

						# Allow selecting many if "multiple select"
						if qtype == "multiple select":
							field["properties"]["allow_multiple_selection"] = True
						else:
							field["properties"]["allow_multiple_selection"] = False
						field["properties"]["randomize"] = True

						fields.append(field)

					if st.session_state.quiz_title_nt.lower().endswith(
						"quiz"
					) or st.session_state.quiz_title_nt.lower().endswith("form"):
						welcome_text = (
							f"Welcome to the {st.session_state.quiz_title_nt}!"
						)
					else:
						welcome_text = (
							f"Welcome to the {st.session_state.quiz_title_nt} Quiz!"
						)

					tys_text = translate_text("Thank you for completing the form!")

					# Creating the full quiz JSON payload
					quiz_payload = {
						"title": st.session_state.quiz_title,
						"type": "quiz",
						"settings": {
							"language": "en",
							"progress_bar": "percentage",
							"autosave_progress": True,
							"is_public": True,
							"meta": {"allow_indexing": False},
							# "mode": "knowledge_quiz"
						},
						"welcome_screens": [
							{
								"title": translate_text(welcome_text),
								"properties": {
									"show_button": True,
									"button_text": "Start Quiz",
								},
							}
						],
						"fields": fields,
						"thankyou_screens": [
							{
								"ref": "static_tys",
								"title": tys_text,
								"type": "thankyou_screen",
								"properties": {
									"show_button": False,
									"share_icons": False,
								},
							}
						],
						"theme": {"href": "https://api.typeform.com/themes/NqJgJG"},
					}

					# Step 4.4: Sending the request to create the form
					form_response = requests.post(
						"https://api.typeform.com/forms",
						headers=headers,
						json=quiz_payload,
					)

					quiz_end_time = time.time()
					try:
						form_json = form_response.json()
						typeform_quiz_url = form_json["_links"]["display"]

					except:
						st.session_state.quiz_error = True
						st.session_state.history_quiz_links.append("failed")
						st.session_state.history_quiz_qrs.append("failed")
						st.session_state.history_quiz_docs.append("failed")

						st.rerun()

					# Step 4.5: Creating the quiz Word document

					# Create the document
					quiz_docx = Document()

					# 1. Name and Class (Left-aligned, Bold, Size 12)
					for label in ["Name", "Class"]:
						p = quiz_docx.add_paragraph()
						run = p.add_run(f"{label}: {"_" * 40}")
						run.font.bold = True
						run.font.size = Pt(12)
						p.alignment = WD_ALIGN_PARAGRAPH.LEFT

					quiz_docx.add_paragraph("")  # spacer

					# 2. Quiz Title (Centered, Bold, Size 30)
					title = quiz_docx.add_paragraph()
					run = title.add_run(f"{st.session_state.quiz_title.strip()} Quiz")
					run.font.bold = True
					run.font.size = Pt(30)
					title.alignment = WD_ALIGN_PARAGRAPH.CENTER

					quiz_docx.add_paragraph("")  # spacer

					# 3. Questions
					for qtext, qchoices_dict, qnum, qtype in zip(
						questions_texts,
						questions_choices,
						range(len(questions_texts)),
						questions_types,
					):
						qnum += 1
						choices_list = list(qchoices_dict.values())
						choices_letters = [
							"A",
							"B",
							"C",
							"D",
							"E",
							"F",
							"G",
							"H",
							"I",
							"J",
						]

						# Question (Bold, size 18)
						question_line = f"Q{qnum}: {qtext}"
						if qtype.strip().lower() == "multiple select":
							question_line += " (multiple select)"

						p = quiz_docx.add_paragraph()
						run = p.add_run(question_line)
						run.font.bold = True
						run.font.size = Pt(18)

						# Choices (Normal, size 18)
						for letter, choice in zip(
							choices_letters[: len(choices_list[0])], choices_list[0]
						):
							p = quiz_docx.add_paragraph()
							run = p.add_run(f"{letter}. {choice}")
							run.font.size = Pt(18)

						quiz_docx.add_paragraph("")  # spacer

					# 4. Page break
					quiz_docx.add_page_break()

					# 5. Model Answer Header
					p = quiz_docx.add_paragraph()
					run = p.add_run("Model Answer")
					run.font.bold = True
					run.font.size = Pt(24)
					run.font.color.rgb = RGBColor(255, 0, 0)
					p.alignment = WD_ALIGN_PARAGRAPH.CENTER

					quiz_docx.add_paragraph("")  # spacer

					# 6. Model Answers List
					for correct_letters_lst, qnum in zip(
						correct_choices, range(len(questions_texts))
					):
						p = quiz_docx.add_paragraph()
						run = p.add_run(f"{qnum + 1}. {', '.join(correct_letters_lst)}")
						run.font.size = Pt(18)

					# 7. Save to buffer
					quiz_docx_buffer = BytesIO()
					quiz_docx.save(quiz_docx_buffer)
					quiz_docx_buffer.seek(0)

					# Step 4.5: Generating the QR code
					quiz_qr = qrcode.make(typeform_quiz_url)
					quiz_qr_buffer = BytesIO()
					quiz_qr.save(quiz_qr_buffer, format="PNG")
					quiz_qr_buffer.seek(0)

					# Step 4.6: Saving reults
					st.session_state.typeform_quiz_url = typeform_quiz_url
					st.session_state.quiz_qr_buffer = quiz_qr_buffer
					st.session_state.quiz_doc_buffer = quiz_docx_buffer
					st.session_state.quiz_time = quiz_end_time - quiz_start_time

					st.session_state.history_quiz_links.append(
						st.session_state.typeform_quiz_url
					)
					st.session_state.history_quiz_qrs.append(
						st.session_state.quiz_qr_buffer
					)
					st.session_state.history_quiz_docs.append(
						st.session_state.quiz_doc_buffer
					)

					st.session_state.quiz_done = True

					with quiz_tab:
						st.write(
							f"➡️  "
							+ translate_text(
								f"Quiz link: {st.session_state.typeform_quiz_url}"
							)
						)

						st.image(
							st.session_state.quiz_qr_buffer,
							caption=translate_text("Scan to take the quiz"),
							width=250,
						)
						st.divider()
						st.download_button(
							label=translate_text("Download Quiz Word Document"),
							data=st.session_state.quiz_doc_buffer,
							file_name=f"{st.session_state.quiz_title_nt.strip()}_quiz.docx",
							mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
						)
						st.divider()
						st.write(
							translate_text(
								f"Made in {st.session_state.quiz_time:.1f} seconds"
							)
						)

			elif st.session_state.quiz_done and not st.session_state.quiz_error:
				with quiz_tab:
					st.write(
						f"➡️  "
						+ translate_text(
							f"Quiz link: {st.session_state.typeform_quiz_url}"
						)
					)

					st.image(
						st.session_state.quiz_qr_buffer,
						caption=translate_text("Scan to take the quiz"),
						width=250,
					)
					st.divider()
					st.download_button(
						label=translate_text("Download Quiz Word Document"),
						data=st.session_state.quiz_doc_buffer,
						file_name=f"{st.session_state.quiz_title.strip()}_quiz.docx",
						mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
					)
					st.divider()
					st.write(
						translate_text(
							f"Made in {st.session_state.quiz_time:.1f} seconds"
						)
					)

			" "
			" "
			st.session_state.script_time = (
				st.session_state.echo3d_time
				+ st.session_state.telegraph_time
				+ st.session_state.quiz_time
			)
			st.write(
				translate_text(
					f"**All made in {st.session_state.script_time:.1f} seconds**"
				)
			)

			st.session_state.backend_running = False

			if st.session_state.history_rerun2:
				st.session_state.history_rerun2 = False
				if st.session_state.ar_skip and st.session_state.ar_error:
					st.session_state.history_ar_links.append("failed")
					st.session_state.history_ar_qrs.append("failed")
				elif st.session_state.ar_skip:
					st.session_state.history_ar_links.append("abandoned")
					st.session_state.history_ar_qrs.append("abandoned")
				st.rerun()

			if st.button("Generate New"):
				st.session_state.model_name = None
				st.session_state.age = None
				st.session_state.model_uid = None
				st.session_state.user_content_length = None
				st.session_state.user_content_complexity = None
				st.session_state.user_custom_prompt = None

				st.session_state.ar_skip = None
				st.session_state.ar_done = False
				st.session_state.content_done = False
				st.session_state.quiz_done = False

				st.session_state.echo3d_time = 0
				st.session_state.telegraph_time = 0
				st.session_state.quiz_time = 0

				st.session_state.ar_error = False
				st.session_state.content_error = False
				st.session_state.quiz_error = False

				st.session_state.backend_running = False
				st.session_state.history_rerun = True
				st.session_state.history_rerun2 = True

				st.rerun()

	# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------

	# --- Streamlit UI Code ---

	# What is teachAR?
	st.header(translate_text("What is TeachAR?"))

	wts_1 = "TeachAR is an educational platform that brings learning to life using **Augmented Reality (AR)**. 🌟 Just **type the name** of any science, health, or engineering concept or **upload an image** and we will provide you:"
	wts_2 = """
- 🧠 **An interactive AR experience** on your mobile or laptop in your real environment.
- 📚 **Fun and clear educational content** written for your age and grade where you can study and learn.  
- 📄 **A Word Document** containing that content.  
- ✅ **A quiz to test what you've learned** — based on the content provided to you. This is perfect for revision or school projects.  
- 📜 **A Word Document** containing the quiz questions and a model answer.
"""
	wts_3 = "TeachAR is designed to make learning **exciting, visual, interactive and unforgettable**! 🎯✨"

	st.write(translate_text(wts_1))
	st.write(translate_text(wts_2))
	st.write(translate_text(wts_3))

	st.divider()
	st.header(translate_text("Special Templates " + "🚀"))

	st.write(
		f"**{translate_text("Try these ready-made examples by TeachAR to see it in action!")}**"
	)

	# Row of 3 cards

	col1, col2, col3 = st.columns(3)
	with col1:
		st.image(
			"assets/images/ST_iphone_16.jpeg",
			caption=translate_text("iPhone 16 Pro Max"),
		)
		st.markdown(
			f"""
			[{translate_text("View in AR")}](https://go.echo3d.co/kb1k)  
			[{translate_text("Read Content")}](https://telegra.ph/Iphone-16-for-Age-13-06-30)  
			[{translate_text("Take Quiz")}](https://form.typeform.com/to/j6PrboSn)
			"""
		)

	with col2:
		st.image(
			"assets/images/ST_human_heart.png", caption=translate_text("Human Heart")
		)
		st.markdown(
			f"""
			[{translate_text("View in AR")}](https://go.echo3d.co/Gq3D)  
			[{translate_text("Read Content")}](https://telegra.ph/Human-Heart-for-Age-24-07-01)  
			[{translate_text("Take Quiz")}](https://form.typeform.com/to/tn4BcNl0)
			"""
		)

	with col3:
		st.image(
			"assets/images/ST_fighter_jet.jpeg", caption=translate_text("Fighter Jet")
		)
		st.markdown(
			f"""
			[{translate_text("View in AR")}](https://go.echo3d.co/M3Pm)  
			[{translate_text("Read Content")}](https://telegra.ph/Fighter-Jet-for-Age-10-06-28)  
			[{translate_text("Take Quiz")}](https://form.typeform.com/to/dWFhzNtu)
			"""
		)

	# Get started
	st.divider()
	st.header(translate_text("Get Started") + " 🚀")

	# Function to get model name from uploaded image bytes
	def get_model_name_from_image(image_bytes):
		# Search by image function
		img = PIL.Image.open(io.BytesIO(image_bytes))

		response = st.session_state.gemini_2_0_flash.generate_content(
			[
				"What is this object in the image? If this object conflicts with another object, try to make it more specific. For example, if it is an organ, define for which organism it is. Strict Notes: You must reply with the only the final answer (the object's name), no explaination or more words. Here is the image:",
				img,
			]
		)
		return response.text.strip()

	if "model_name" not in st.session_state or "age" not in st.session_state:
		# Model Name
		model_name = st.text_input(
			translate_text("Topic"),
			placeholder=translate_text("Enter the name of the model"),
		)
		# Image Upload
		if not model_name:
			st.session_state.is_uploaded = True
			uploaded_image = st.file_uploader(
				"OR upload an image of an object", type=["jpg", "jpeg", "png"]
			)
		# Age
		" "
		age = st.number_input(
			translate_text("Age"), min_value=0, max_value=100, value=10
		)

		# Content Configurations
		" "
		with st.expander(translate_text("Content Configurations")):
			# Grade
			grade = st.text_input(
				translate_text("Grade"),
				placeholder=translate_text(
					"Enter your grade (e.g. second preparatory)"
				),
				value=st.session_state.user_grade,
			)
			# Content Length
			content_length = st.radio(
				label=translate_text("Content Length"),
				options=[
					translate_text("Auto"),
					translate_text("Short"),
					translate_text("Medium"),
					translate_text("Long"),
					translate_text("Very Long"),
				],
				horizontal=True,
				index=0,
			)
			# Content Complexity
			content_complexity = st.radio(
				label=translate_text("Content Complexity"),
				options=[
					translate_text("Auto"),
					translate_text("Easy"),
					translate_text("Medium"),
					translate_text("Hard"),
				],
				horizontal=True,
				index=0,
			)

			# Custom Prompt
			custom_prompt = st.text_area(translate_text("Custom Prompt"))

			" "
			st.info(
				translate_text(
					"Note: The quiz length and info depend on the content length."
				)
			)

		generate_button = st.button(translate_text("Generate"))
		if generate_button and (model_name or uploaded_image) and age:

			if "is_uploaded" in st.session_state:
				if model_name:
					st.session_state.model_name = model_name

				elif uploaded_image:
					image_bytes = uploaded_image.read()
					with st.spinner(translate_text("Identifying the object...")):
						detected_model_name = get_model_name_from_image(image_bytes)
					if detected_model_name:
						st.session_state.model_name = detected_model_name
						st.success(translate_text(f"Object: **{detected_model_name}**"))
					else:
						st.warning(
							translate_text(
								"Could not identify the object. Please try another image."
							)
						)
						st.stop()  # stops further execution if model name failed

				else:
					st.session_state.model_name = model_name

			st.session_state.age = age

			st.session_state.user_grade = grade
			st.session_state.user_content_length = content_length
			st.session_state.user_content_complexity = content_complexity
			st.session_state.user_custom_prompt = custom_prompt

			automated_processes(st.session_state.model_name, st.session_state.age)

		elif generate_button:
			st.warning(translate_text("Please fill in the name and age."))
	else:
		if st.session_state.model_name is None or st.session_state.age is None:
			# Model Name
			model_name = st.text_input(
				translate_text("Topic"),
				placeholder=translate_text("Enter the name of the model"),
			)
			# Image Upload
			if not model_name:
				st.session_state.is_uploaded = True
				uploaded_image = st.file_uploader(
					"OR upload an image of an object", type=["jpg", "jpeg", "png"]
				)
			# Age
			" "
			age = st.number_input(
				translate_text("Age"), min_value=0, max_value=100, value=10
			)

			# Content Configurations
			" "
			with st.expander(translate_text("Content Configurations")):
				# Grade
				grade = st.text_input(
					translate_text("Grade"),
					placeholder=translate_text(
						"Enter your grade (e.g. second preparatory)"
					),
					value=st.session_state.user_grade,
				)
				# Content Length
				content_length = st.radio(
					label=translate_text("Content Length"),
					options=[
						translate_text("Auto"),
						translate_text("Short"),
						translate_text("Medium"),
						translate_text("Long"),
						translate_text("Very Long"),
					],
					horizontal=True,
					index=0,
				)
				# Content Complexity
				content_complexity = st.radio(
					label=translate_text("Content Complexity"),
					options=[
						translate_text("Auto"),
						translate_text("Easy"),
						translate_text("Medium"),
						translate_text("Hard"),
					],
					horizontal=True,
					index=0,
				)

				# Custom Prompt
				custom_prompt = st.text_area(translate_text("Custom Prompt"))

			generate_button = st.button(translate_text("Generate"))
			if generate_button and (model_name or uploaded_image) and age:

				if "is_uploaded" in st.session_state:
					if model_name:
						st.session_state.model_name = model_name

					elif uploaded_image:
						image_bytes = uploaded_image.read()
						with st.spinner(translate_text("Identifying the object...")):
							detected_model_name = get_model_name_from_image(image_bytes)
						if detected_model_name:
							st.session_state.model_name = detected_model_name
							st.success(
								translate_text(f"Object: **{detected_model_name}**")
							)
						else:
							st.warning(
								translate_text(
									"Could not identify the object. Please try another image."
								)
							)
							st.stop()  # stops further execution if model name failed

					else:
						st.session_state.model_name = model_name

				st.session_state.age = age

				st.session_state.user_grade = grade
				st.session_state.user_content_length = content_length
				st.session_state.user_content_complexity = content_complexity
				st.session_state.user_custom_prompt = custom_prompt

				automated_processes(st.session_state.model_name, st.session_state.age)

			elif generate_button:
				st.warning(translate_text("Please fill in the name and age."))

		else:
			model_name = st.session_state.model_name
			age = st.session_state.age
			automated_processes(model_name, age)

# -----------------------------
# 📄 Upload PDF & Summarize
# -----------------------------

elif "upload pdf & summarize" in st.session_state.current_page.lower():
	st.title("🧠 Text & PDF Summarizer")

	summarization_type = st.radio(
		"Choose input method:", ["📜 Enter Text", "📄 Upload PDF"]
	)

	if summarization_type == "📜 Enter Text":
		input_text = st.text_area("✍ Paste or type your text here:", height=300)

		if st.button("🧠 Summarize Text"):
			if input_text.strip() == "":
				st.warning("⚠ Please enter some text to summarize.")
			else:
				with st.spinner("Summarizing your text..."):
					prompt = f"Summarize this clearly for a student:\n\n{input_text}"
					response = st.session_state.gemini_2_0_flash.generate_content(prompt)
					st.markdown("### 📝 Text Summary:")
					st.write(response.text)

	elif summarization_type == "📄 Upload PDF":
		pdf_file = st.file_uploader("📄 Upload a PDF file", type=["pdf"])

		if pdf_file:
			reader = PdfReader(pdf_file)
			full_text = ""
			for page in reader.pages:
				full_text += page.extract_text() or ""

			if not full_text.strip():
				st.error("❌ Could not extract any readable text from the PDF.")
			else:
				if st.button("📚 Summarize PDF"):
					with st.spinner("Summarizing your PDF..."):
						summary_prompt = (
							f"Summarize this clearly for a student:\n\n{full_text}"
						)
						summary_response = (
							st.session_state.gemini_2_0_flash.generate_content(
								summary_prompt
							)
						)
						st.markdown("### 📑 PDF Summary:")
						st.write(summary_response.text)


# -----------------------------
# ❓ Quiz Generator (MCQ + Essay)
# -----------------------------

elif "quiz generator" in st.session_state.current_page.lower():
	st.title("❓ Generate Quiz")

	quiz_topic = st.text_input(
		"Enter a topic for the quiz (e.g. Photosynthesis, Climate Change)"
	)

	num_questions = st.number_input(
		"How many questions?", min_value=1, max_value=30, value=10, step=1
	)

	difficulty = st.selectbox("Select difficulty level:", ["Easy", "Medium", "Hard"])

	question_type = st.selectbox(
		"Select question type:", ["Multiple Choice", "Essay Questions"]
	)

	if "quiz_data" not in st.session_state:
		st.session_state.quiz_data = None
	if "quiz_answers" not in st.session_state:
		st.session_state.quiz_answers = {}

	if st.button("🎯 Generate Quiz"):
		if question_type == "Multiple Choice":
			quiz_prompt = f"""
			You are an educational quiz generator.
			Create {num_questions} multiple-choice questions about "{quiz_topic}".
			Each question must have 4 answer options (A, B, C, D).
			One correct answer must be clearly marked at the end as: Answer: X (where X is A/B/C/D).
			Format:
			1) Question text
			A) Option A
			B) Option B
			C) Option C
			D) Option D
			Answer: X
			Difficulty: {difficulty} level.
			"""
		else:
			quiz_prompt = f"""
			You are an educational quiz generator.
			Create {num_questions} open-ended essay questions about "{quiz_topic}".
			Each question should match a {difficulty} level.
			Format:
			1) Question text
			"""

		quiz_response = st.session_state.gemini_2_0_flash.generate_content(quiz_prompt)
		quiz_text = quiz_response.text

		if question_type == "Multiple Choice":
			questions = []
			blocks = quiz_text.strip().split("\n\n")
			for block in blocks:
				lines = block.strip().split("\n")
				if len(lines) >= 6:
					question_text = lines[0][3:].strip()
					options = {
						"A": lines[1][3:].strip(),
						"B": lines[2][3:].strip(),
						"C": lines[3][3:].strip(),
						"D": lines[4][3:].strip(),
					}
					correct_line = [l for l in lines if l.startswith("Answer:")]
					correct_answer = (
						correct_line[0].split("Answer:")[1].strip()
						if correct_line
						else ""
					)
					questions.append(
						{
							"question": question_text,
							"options": options,
							"answer": correct_answer,
						}
					)
			st.session_state.quiz_data = questions

		else:
			questions = []
			lines = quiz_text.strip().split("\n")
			for line in lines:
				if line.strip() and line[0].isdigit() and ")" in line:
					q_text = line.split(")", 1)[1].strip()
					questions.append({"question": q_text})
			st.session_state.quiz_data = questions

		st.session_state.quiz_answers = {}

	if st.session_state.quiz_data:
		st.subheader("📋 Take the Quiz")

		if question_type == "Multiple Choice":
			for idx, q in enumerate(st.session_state.quiz_data):
				st.markdown(f"Q{idx+1}: {q['question']}")
				selected = st.radio(
					f"Your Answer for Q{idx+1}:",
					list(q["options"].keys()),
					format_func=lambda x: f"{x}) {q['options'][x]}",
					key=f"q_{idx}",
				)
				st.session_state.quiz_answers[idx] = selected

			if st.button("✅ Submit Quiz"):
				score = 0
				total = len(st.session_state.quiz_data)
				results = []
				for idx, q in enumerate(st.session_state.quiz_data):
					user_ans = st.session_state.quiz_answers.get(idx, None)
					correct_ans = q["answer"]
					if user_ans == correct_ans:
						score += 1
					result_line = (
						f"Q{idx+1}: ✅ Correct!"
						if user_ans == correct_ans
						else f"Q{idx+1}: ❌ Incorrect (Your: {user_ans}, Correct: {correct_ans})"
					)
					results.append(result_line)

				st.success(f"Your score: {score} / {total}")
				st.markdown("### 🗒 Detailed Results")
				for res in results:
					st.write(res)

				# SAVE TO HISTORY
				add_chat_to_history(
					st.session_state.user,
					"quiz",
					{
						"topic": quiz_topic,
						"type": "MCQ",
						"score": f"{score}/{total}",
						"results": results,
					},
				)

				# Reset
				st.session_state.quiz_data = None
				st.session_state.quiz_answers = {}

		else:
			for idx, q in enumerate(st.session_state.quiz_data):
				st.markdown(f"Q{idx+1}: {q['question']}")
				essay_ans = st.text_area(
					f"Your Answer for Q{idx+1}:", key=f"essay_{idx}"
				)
				st.session_state.quiz_answers[idx] = essay_ans

			if st.button("✅ Submit Essay Quiz"):
				results = []
				for idx, q in enumerate(st.session_state.quiz_data):
					user_ans = st.session_state.quiz_answers.get(idx, "")
					results.append(
						f"Q{idx+1}: {q['question']}\nYour Answer: {user_ans}\n"
					)

				st.success("✅ Your essay answers have been saved to your history.")
				st.markdown("### 🗒 Your Answers")
				for res in results:
					st.write(res)

				# SAVE TO HISTORY
				add_chat_to_history(
					st.session_state.user,
					"quiz",
					{"topic": quiz_topic, "type": "Essay", "results": results},
				)

				# Reset
				st.session_state.quiz_data = None
				st.session_state.quiz_answers = {}

# -----------------------------
# 📜 History
# -----------------------------

elif (
	not st.session_state.backend_running
	and "history" in st.session_state.current_page.lower()
):

	st.title("📂 " + translate_text("History Page"))
	" "

	if st.session_state.history_model_names:
		for index, (
			model_name,
			model_thumb,
			ar_link,
			ar_qr,
			content_link,
			content_qr,
			content_doc,
			quiz_link,
			quiz_qr,
			quiz_doc,
		) in enumerate(
			zip(
				st.session_state.history_model_names,
				st.session_state.history_model_thumbs,
				st.session_state.history_ar_links,
				st.session_state.history_ar_qrs,
				st.session_state.history_content_links,
				st.session_state.history_content_qrs,
				st.session_state.history_content_docs,
				st.session_state.history_quiz_links,
				st.session_state.history_quiz_qrs,
				st.session_state.history_quiz_docs,
			)
		):
			st.markdown(f"### 🧠 {translate_text(f"Model {index + 1}: {model_name}")}")

			# Ensure model_thumb is a valid URL or image buffer
			if isinstance(model_thumb, str) and model_thumb.startswith("http"):
				st.image(model_thumb, caption=model_name, width=300)

			h_ar_tab, h_content_tab, h_quiz_tab = st.tabs(
				["AR", translate_text("Content"), translate_text("Quiz")]
			)

			with st.expander("View Details"):
				with h_ar_tab:
					# Check if ar_link is 'abandoned' or a real link
					if ar_link == "abandoned":
						st.info(translate_text("You abandoned the AR for this item."))
					elif ar_link == "failed":
						st.error(translate_text("AR generation failed for this item."))
					else:
						st.write(f"➡️  **{translate_text("AR Link")}:** {ar_link}")
						ar_qr.seek(0)
						st.image(
							ar_qr,
							caption=translate_text("Scan with your phone"),
							width=250,
						)

				with h_content_tab:
					if content_link != "failed":
						st.write(
							f"➡️  **{translate_text("Content Link")}:** {content_link}"
						)

						if isinstance(content_qr, BytesIO):
							content_qr.seek(0)
						st.image(
							content_qr,
							caption=translate_text("Scan to open the content"),
							width=250,
						)

						st.divider()

						if isinstance(content_doc, BytesIO):
							content_doc.seek(0)

						st.download_button(
							label="📄 " + translate_text("Download Content Document"),
							data=content_doc,
							file_name=f"{model_name}_content.docx",
							mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
							key=f"content_download_{index}",
						)
					else:
						st.error(
							translate_text("Content generation failed for this item.")
						)

				with h_quiz_tab:
					if quiz_link != "failed":
						st.write(f"➡️  **{translate_text("Quiz Link")}:** {quiz_link}")

						if isinstance(quiz_qr, BytesIO):
							quiz_qr.seek(0)
						st.image(
							quiz_qr,
							caption=translate_text("Scan to take the quiz"),
							width=250,
						)

						st.divider()

						if isinstance(quiz_doc, BytesIO):
							quiz_doc.seek(0)

						st.download_button(
							label="📄 " + translate_text("Download Quiz Word Document"),
							data=quiz_doc,
							file_name=f"{model_name}_quiz.docx",
							mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
							key=f"quiz_download_{index}",
						)
					else:
						st.error(
							translate_text("Quiz generation failed for this item.")
						)

			" "
			" "
			st.divider()
			" "
			" "
	else:
		st.info(translate_text("No history yet. Generate a model to see it here."))


